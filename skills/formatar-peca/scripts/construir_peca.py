#!/usr/bin/env python3
"""
construir_peca.py — Gerador nativo RDAA (python-docx puro, sem docxtpl)

Substitui o fluxo antigo de "template master + docxtpl" para o corpo da peça.
Motivo: docxtpl injeta cada campo como uma única string (via {{ campo }}), o
que torna impossível aplicar formatação por trecho dentro de um mesmo
parágrafo (ex.: nome da parte em negrito+sublinhado só na abertura, título
em negrito com tabulação exata, alínea com recuo deslocado). Este script
constrói o documento inteiro programaticamente, parágrafo a parágrafo, com
a formatação correta aplicada em cada tipo de bloco.

Todos os valores de recuo/tabulação abaixo foram conferidos byte a byte
contra uma peça real já revisada e aprovada do escritório (conversão para
twips: 1 cm = 567, 2 cm = 1134, 3 cm = 1701).

Uso:
    python3 construir_peca.py --context /tmp/rdaa_context.json --output outputs/peca.docx

Formato do JSON de contexto — ver references/schema_blocos.md.
"""

import argparse
import json
import os
import re
import sys
import tempfile

# ponytail: no Windows, stderr/stdout herdam o codepage ANSI do console (ex.:
# cp1252) quando o processo é aberto via subprocess sem PYTHONUTF8/PYTHONIOENCODING
# — uma mensagem com acento (ex.: "campo obrigatório") sai como bytes UTF-8
# decodificados errado do outro lado ("campo obrigatÃ³rio"). Reconfigurar aqui
# garante saída correta não importa quem chame o script (pytest, Bash tool,
# executar_motor.py).
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(encoding="utf-8")

try:
    from docx import Document
    from docx.shared import Pt, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx',
                           '--break-system-packages', '-q'])
    from docx import Document
    from docx.shared import Pt, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ── Constantes RDAA ───────────────────────────────────────────────────────────

FONTE     = 'Tahoma'
TAMANHO   = Pt(10.5)
COR       = RGBColor(0, 0, 0)
COR_DOURADA = RGBColor(0xFF, 0xC0, 0x00)  # padrão da peça real assinada, linha do site no rodapé
TITULO_PROIBIDO_CHARS = {'—': 'travessão', '–': 'travessão', ':': 'dois pontos'}

# Endereço em title-case (não caixa alta) — Apontamentos de melhoria 2026-07,
# alinhado ao padrão da peça real corrigida.
RODAPE_L1 = (
    'Avenida dos Vinhedos, 200 | Conj. 4 | Gavea Office | '
    'Morada da Colina | 38411-159 | Uberlândia/MG | +55 34 3229 0020'
)
RODAPE_L2 = 'romanodonadel.com.br'

PUBLICACOES_PADRAO = (
    'Requer que as publicações referentes a este feito sejam realizadas exclusivamente '
    'em nome do advogado WANDERLEY ROMANO DONADEL, OAB/MG 78.870, '
    'wanderley@romanodonadel.com.br, e as eventualmente postais encaminhadas ao '
    'seguinte endereço, Avenida dos Vinhedos, n.º 200, Conjunto 4, Gavea Office, '
    'Morada da Colina, Uberlândia/MG, CEP 38.411-159, sob pena de nulidade.'
)

# Grade fixa 2x2 (ordem = posição na tabela: [0]=linha1/col1, [1]=linha1/col2,
# [2]=linha2/col1, [3]=linha2/col2) — confirmado com o Ricardo em 2026-07-19
# contra peça real assinada. Os 4 advogados aparecem sempre, nesta ordem.
SIGNATARIOS = [
    ('Wanderley Romano Donadel',              'OAB/MG 78.870',  'wanderley@romanodonadel.com.br',            'Assinado Eletronicamente'),
    ('Flávia Almeida Forti da Fonseca',       'OAB/MG 96.919',  'flavia@romanodonadel.com.br',                None),
    ('Alessandra Xavier Coutinho',            'OAB/MG 208.101', 'alessandra.coutinho@romanodonadel.com.br',  None),
    ('Ricardo Cesar Souza de Oliveira Junior', 'OAB/MG 208.090', 'ricardo.oliveira@romanodonadel.com.br',    None),
]

# Alternância explícita por UF do processo originário. O gerador não tenta
# descobrir a UF pelo número, endereçamento ou texto livre.
OAB_WANDERLEY_POR_UF = {
    'MG': 'OAB/MG 78.870', 'GO': 'OAB/GO 18.703-A', 'RJ': 'OAB/RJ 223.820-A',
    'SP': 'OAB/SP 422.887-A', 'MA': 'OAB/MA 28.946-A', 'SC': 'OAB/SC 73.406',
    'AM': 'OAB/AM 1.660-A', 'RO': 'OAB/RO 12.856-A', 'ES': 'OAB/ES 38.504-A',
    'MS': 'OAB/MS 29.807-A', 'BA': 'OAB/BA 83.986', 'PE': 'OAB/PE 67.943',
}
OAB_FLAVIA_POR_UF = {'MG': 'OAB/MG 96.919', 'SP': 'OAB/SP 548.105'}
UF_NOMES = {
    'MINAS GERAIS': 'MG', 'GOIAS': 'GO', 'GOIÁS': 'GO', 'RIO DE JANEIRO': 'RJ',
    'SAO PAULO': 'SP', 'SÃO PAULO': 'SP', 'MARANHAO': 'MA', 'MARANHÃO': 'MA',
    'SANTA CATARINA': 'SC', 'AMAZONAS': 'AM', 'RONDONIA': 'RO', 'RONDÔNIA': 'RO',
    'ESPIRITO SANTO': 'ES', 'ESPÍRITO SANTO': 'ES', 'MATO GROSSO DO SUL': 'MS',
    'BAHIA': 'BA', 'PERNAMBUCO': 'PE',
}
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)


def TW(twips):
    """Converte twips exatos para EMU (1 twip = 635 EMU). Evita o arredondamento
    de Cm()/Mm() para os valores 567/1134/1701 exigidos pela especificação."""
    return Emu(round(twips * 635))


CM1 = TW(567)    # 1 cm
CM2 = TW(1134)   # 2 cm
CM3 = TW(1701)   # 3 cm


# ── Infraestrutura de numeração nativa ────────────────────────────────────
#
# Migra títulos (I., II., III.), parágrafos numerados (1., 2., 3.) e alíneas
# (a), b), c)) de texto digitado (f'{numeral}\t') para listas reais do Word
# (w:abstractNum/w:num/w:numPr). Corrige o bug real do "10." duplicado
# (causa raiz: quem redige perde a conta manualmente) e permite renumeração
# automática no Word após edição.


def _criar_elemento_lvl(ilvl, num_fmt, lvl_text, start_val, left_twips, hanging_twips, tab_twips, jc='left', bold=False):
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), str(ilvl))

    start = OxmlElement('w:start')
    start.set(qn('w:val'), str(start_val))
    lvl.append(start)

    fmt = OxmlElement('w:numFmt')
    fmt.set(qn('w:val'), num_fmt)
    lvl.append(fmt)

    lt = OxmlElement('w:lvlText')
    lt.set(qn('w:val'), lvl_text)
    lvl.append(lt)

    ljc = OxmlElement('w:lvlJc')
    ljc.set(qn('w:val'), jc)
    lvl.append(ljc)

    suff = OxmlElement('w:suff')
    suff.set(qn('w:val'), 'tab')
    lvl.append(suff)

    pPr = OxmlElement('w:pPr')
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'num')
    tab.set(qn('w:pos'), str(tab_twips))
    tabs.append(tab)
    pPr.append(tabs)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left_twips))
    ind.set(qn('w:hanging'), str(hanging_twips))
    pPr.append(ind)
    lvl.append(pPr)

    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONTE)
    rFonts.set(qn('w:hAnsi'), FONTE)
    rFonts.set(qn('w:cs'), FONTE)
    rPr.append(rFonts)
    if bold:
        b_el = OxmlElement('w:b')
        rPr.append(b_el)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 10.5pt em meio-pontos
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    lvl.append(rPr)

    return lvl


def _registrar_lista(ct_numbering, num_fmt, lvl_text, start_val,
                     left_twips, hanging_twips, tab_twips, jc='left', bold=False,
                     lvl1_spec=None):
    """Cria um w:abstractNum (nível 0 e opcionalmente nível 1) e insere antes do
    primeiro w:num existente. Retorna o abstractNumId atribuído."""
    existing = ct_numbering.findall(qn('w:abstractNum'))
    max_id = max((int(a.get(qn('w:abstractNumId'))) for a in existing), default=-1)
    new_id = max_id + 1

    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), str(new_id))

    mlt = OxmlElement('w:multiLevelType')
    mlt.set(qn('w:val'), 'multilevel' if lvl1_spec else 'singleLevel')
    abstractNum.append(mlt)

    lvl0 = _criar_elemento_lvl(0, num_fmt, lvl_text, start_val, left_twips, hanging_twips, tab_twips, jc, bold)
    abstractNum.append(lvl0)

    if lvl1_spec:
        # lvl1_spec é dict com chaves: num_fmt, lvl_text, start_val, left_twips, hanging_twips, tab_twips
        lvl1 = _criar_elemento_lvl(1, lvl1_spec['num_fmt'], lvl1_spec['lvl_text'],
                                  lvl1_spec.get('start_val', 1), lvl1_spec['left_twips'],
                                  lvl1_spec['hanging_twips'], lvl1_spec['tab_twips'],
                                  lvl1_spec.get('jc', 'left'), lvl1_spec.get('bold', False))
        abstractNum.append(lvl1)

    first_num = ct_numbering.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstractNum)
    else:
        ct_numbering.append(abstractNum)

    return new_id


def _nova_sequencia(ct_numbering, abstract_num_id):
    """Cria nova sequência (w:num) referenciando o abstractNum dado.
    Cada chamada produz um numId independente que reinicia a contagem."""
    num = ct_numbering.add_num(abstract_num_id)
    return int(num.get(qn('w:numId')))


def _aplicar_numeracao(paragraph, num_id, ilvl=0):
    """Seta w:numPr (ilvl + numId) no parágrafo — ativa a numeração nativa."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)
    pPr.append(numPr)


def _semantic_ids_from_block(block):
    """Retorna IDs declarados no bloco, preservando a ordem e sem inferência."""
    ids = []
    for key in ('id', 'semantic_ids', 'fact_ids', 'thesis_ids', 'request_ids', 'source_ids', 'risk_ids'):
        value = block.get(key)
        values = value if isinstance(value, list) else [value]
        for item in values:
            if item is None:
                continue
            text = str(item).strip()
            if text and text not in ids:
                ids.append(text)
    return ids


def _semantic_bookmark_name(semantic_id, occurrence):
    safe = re.sub(r'[^A-Za-z0-9-]+', '-', str(semantic_id)).strip('-') or 'sem-id'
    return f'rdaa_{safe}__{occurrence}'


def _paragraph_for_semantic_result(result):
    if result is None:
        return None
    values = result if isinstance(result, list) else [result]
    for value in values:
        if hasattr(value, '_p'):
            return value._p
        if hasattr(value, '_tbl'):
            try:
                return value.rows[0].cells[0].paragraphs[0]._p
            except (IndexError, AttributeError):
                continue
    return None


def _mark_semantic_ids(result, block, occurrences):
    """Insere bookmarks zero-width, invisíveis, no primeiro parágrafo gerado."""
    paragraph = _paragraph_for_semantic_result(result)
    if paragraph is None:
        return
    for semantic_id in _semantic_ids_from_block(block):
        occurrence = int(occurrences.get(semantic_id, 0)) + 1
        occurrences[semantic_id] = occurrence
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(10000 + sum(occurrences.values())))
        bookmark_start.set(qn('w:name'), _semantic_bookmark_name(semantic_id, occurrence))
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_start.get(qn('w:id')))
        ppr = paragraph.find(qn('w:pPr'))
        insert_at = 1 if ppr is not None else 0
        paragraph.insert(insert_at, bookmark_start)
        paragraph.insert(insert_at + 1, bookmark_end)


def _preparar_numeracao(doc):
    """Registra as listas base no numbering.xml do documento e retorna ({tipo: abstractNumId}, ct_numbering)."""
    ct = doc.part.numbering_part.numbering_definitions._numbering

    # Título Nível 1: I., II., III. (2cm / tab 2cm, negrito)
    titulo_id = _registrar_lista(
        ct, 'upperRoman', '%1.', 1,
        left_twips=1134, hanging_twips=1134, tab_twips=1134,
        bold=True)

    # Título Nível 2: 1., 2., 3. (centralizado, negrito, só 1ª letra maiúscula)
    titulo2_id = _registrar_lista(
        ct, 'decimal', '%1.', 1,
        left_twips=0, hanging_twips=0, tab_twips=0,
        jc='center', bold=True)

    # Título Nível 3: a., b., c. (4cm / tab 4cm, borda inferior, CAIXA ALTA, negrito)
    titulo3_id = _registrar_lista(
        ct, 'lowerLetter', '%1.', 1,
        left_twips=2268, hanging_twips=1134, tab_twips=2268,
        bold=True)

    # Parágrafo Numerado: 1., 2., 3. (Numeral na margem 0, tab 2cm em 1134, 2ª linha volta à margem 0 - Correções.md §5.1)
    numerado_id = _registrar_lista(
        ct, 'decimal', '%1.', 1,
        left_twips=0, hanging_twips=0, tab_twips=1134)

    # Alíneas multinível: nível 0 = a), b)... (3cm); nível 1 = i), ii)... (4cm)
    alinea_id = _registrar_lista(
        ct, 'lowerLetter', '%1)', 1,
        left_twips=1701, hanging_twips=567, tab_twips=1701,
        lvl1_spec={
            'num_fmt': 'lowerRoman', 'lvl_text': '%2)', 'start_val': 1,
            'left_twips': 2268, 'hanging_twips': 567, 'tab_twips': 2268
        })

    # Lista de Documentos: 1), 2), 3) (3cm, negrito)
    documento_id = _registrar_lista(
        ct, 'decimal', '%1)', 1,
        left_twips=1701, hanging_twips=567, tab_twips=1701,
        bold=True)

    return {
        'titulo': titulo_id,
        'titulo2': titulo2_id,
        'titulo3': titulo3_id,
        'numerado': numerado_id,
        'alinea': alinea_id,
        'documento': documento_id
    }, ct


class _GerenciadorSequencias:
    """Mantém {(tipo_lista, nome_sequencia): num_id} — cada sequência é uma
    contagem independente dentro do mesmo tipo de lista (ex.: parágrafos
    numerados reiniciando entre petição de interposição e razões)."""

    def __init__(self, ct_numbering, abstract_ids):
        self._ct = ct_numbering
        self._abstract_ids = abstract_ids
        self._sequencias = {}

    def num_id_para(self, tipo_lista, sequencia='default', reiniciar=False):
        """Retorna num_id existente ou cria nova sequência se reiniciar=True
        ou se a sequência é inédita."""
        chave = (tipo_lista, sequencia)
        if reiniciar or chave not in self._sequencias:
            num_id = _nova_sequencia(self._ct, self._abstract_ids[tipo_lista])
            self._sequencias[chave] = num_id
        return self._sequencias[chave]


# ── Estilos Nomeados e Inline Formatting (Fase 4) ──────────────────────────

def _adicionar_texto_formatado(paragraph, texto, bold_default=False, italic_default=False, size=TAMANHO):
    """Adiciona runs ao parágrafo interpretando marcações inline simples:
    <b>...</b> ou **...** -> negrito
    <i>...</i> ou *...*   -> itálico
    Se não houver tag inline, adiciona um único run com os padrões passados."""
    if not isinstance(texto, str) or ('<' not in texto and '*' not in texto):
        r = paragraph.add_run(texto)
        _fmt_run(r, bold=bold_default, italic=italic_default, size=size)
        return [r]

    # Regex para capturar tags HTML básicas (<b>, <i>) e markdown (** e *)
    pattern = r'(<b>.*?</b>|<i>.*?</i>|\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, texto, flags=re.DOTALL)
    runs = []

    for part in parts:
        if not part:
            continue
        is_b = bold_default
        is_i = italic_default
        sub_text = part

        if part.startswith('<b>') and part.endswith('</b>'):
            is_b = True
            sub_text = part[3:-4]
        elif part.startswith('<i>') and part.endswith('</i>'):
            is_i = True
            sub_text = part[3:-4]
        elif part.startswith('**') and part.endswith('**'):
            is_b = True
            sub_text = part[2:-2]
        elif part.startswith('*') and part.endswith('*'):
            is_i = True
            sub_text = part[1:-1]

        r = paragraph.add_run(sub_text)
        _fmt_run(r, bold=is_b, italic=is_i, size=size)
        runs.append(r)

    return runs


def _criar_estilos_rdaa(doc):
    """Configura os estilos de parágrafo RDAA no documento para que fiquem
    visíveis no painel de navegação do Word e garantam consistência visual."""
    styles = doc.styles

    def get_or_create(name, base_name='Normal'):
        try:
            return styles[name]
        except KeyError:
            return styles.add_style(name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH

    st_t1 = get_or_create('RDAA Título 1')
    st_t1.font.name = FONTE
    st_t1.font.size = TAMANHO
    st_t1.font.bold = True
    st_t1.font.color.rgb = COR

    st_t2 = get_or_create('RDAA Título 2')
    st_t2.font.name = FONTE
    st_t2.font.size = TAMANHO
    st_t2.font.bold = True
    st_t2.font.color.rgb = COR

    st_t3 = get_or_create('RDAA Título 3')
    st_t3.font.name = FONTE
    st_t3.font.size = TAMANHO
    st_t3.font.bold = True
    st_t3.font.color.rgb = COR

    st_tr = get_or_create('RDAA Título Razões')
    st_tr.font.name = FONTE
    st_tr.font.size = TAMANHO
    st_tr.font.bold = True
    st_tr.font.color.rgb = COR

    st_num = get_or_create('RDAA Numerado')
    st_num.font.name = FONTE
    st_num.font.size = TAMANHO
    st_num.font.color.rgb = COR

    st_ali = get_or_create('RDAA Alínea')
    st_ali.font.name = FONTE
    st_ali.font.size = TAMANHO
    st_ali.font.color.rgb = COR

    st_doc = get_or_create('RDAA Documento')
    st_doc.font.name = FONTE
    st_doc.font.size = TAMANHO
    st_doc.font.color.rgb = COR

    st_cit = get_or_create('RDAA Citação')
    st_cit.font.name = FONTE
    st_cit.font.size = Pt(9)
    st_cit.font.color.rgb = COR


# ── Infraestrutura de Notas de Rodapé Nativas (Fase 7) ─────────────────────

class _GerenciadorNotasRodape:
    """Gerencia a criação e injeção de notas de rodapé reais (word/footnotes.xml)
    no pacote OPC do documento do Word."""

    def __init__(self, doc):
        self.doc = doc
        self.notas = []  # [(id, texto_nota)]

    def adicionar_nota(self, paragraph, texto_nota):
        """Adiciona uma nota de rodapé ao final do parágrafo:
        - Injeta <w:footnoteReference w:id="N"/> no parágrafo.
        - Registra a nota para ser gravada no word/footnotes.xml no salvamento."""
        nota_id = len(self.notas) + 1
        self.notas.append((nota_id, texto_nota))

        # Run com a referência sobrescrita no texto do parágrafo
        r = paragraph.add_run()
        rPr = r._r.get_or_add_rPr()
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'FootnoteReference')
        rPr.append(rStyle)

        fnRef = OxmlElement('w:footnoteReference')
        fnRef.set(qn('w:id'), str(nota_id))
        r._r.append(fnRef)
        return nota_id

    def finalizar(self):
        """Se houver notas, cria word/footnotes.xml e adiciona a relação OPC."""
        if not self.notas:
            return

        part_doc = self.doc.part
        package = part_doc.package

        # Verificar se word/footnotes.xml já existe
        footnotes_part = None
        for part in package.parts:
            if part.partname == '/word/footnotes.xml':
                footnotes_part = part
                break

        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        NS_MAP = {'w': W}

        if footnotes_part is None:
            # Criar elemento raiz <w:footnotes>
            footnotes_elm = OxmlElement('w:footnotes')

            # Nota id -1 (separator)
            fn_sep = OxmlElement('w:footnote')
            fn_sep.set(qn('w:type'), 'separator')
            fn_sep.set(qn('w:id'), '-1')
            p_sep = OxmlElement('w:p')
            r_sep = OxmlElement('w:r')
            r_sep.append(OxmlElement('w:separator'))
            p_sep.append(r_sep)
            fn_sep.append(p_sep)
            footnotes_elm.append(fn_sep)

            # Nota id 0 (continuationSeparator)
            fn_csep = OxmlElement('w:footnote')
            fn_csep.set(qn('w:type'), 'continuationSeparator')
            fn_csep.set(qn('w:id'), '0')
            p_csep = OxmlElement('w:p')
            r_csep = OxmlElement('w:r')
            r_csep.append(OxmlElement('w:continuationSeparator'))
            p_csep.append(r_csep)
            fn_csep.append(p_csep)
            footnotes_elm.append(fn_csep)

            # Adicionar notas reais do usuário
            for nid, texto in self.notas:
                fn = OxmlElement('w:footnote')
                fn.set(qn('w:id'), str(nid))
                p = OxmlElement('w:p')

                # Marcação de referência da nota
                r_ref = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), 'FootnoteReference')
                rPr.append(rStyle)
                r_ref.append(rPr)
                r_ref.append(OxmlElement('w:footnoteRef'))
                p.append(r_ref)

                # Texto da nota
                r_txt = OxmlElement('w:r')
                rPr_txt = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Tahoma')
                rFonts.set(qn('w:hAnsi'), 'Tahoma')
                rPr_txt.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '18')  # 9pt
                rPr_txt.append(sz)
                r_txt.append(rPr_txt)
                t = OxmlElement('w:t')
                t.text = f" {texto}"
                r_txt.append(t)
                p.append(r_txt)

                fn.append(p)
                footnotes_elm.append(fn)

            # Criar a nova part no pacote OPC
            from docx.opc.packuri import PackURI
            from docx.opc.part import XmlPart

            reltype = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
            partname = PackURI('/word/footnotes.xml')
            footnotes_part = XmlPart(partname, 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml', footnotes_elm, package)
            part_doc.relate_to(footnotes_part, reltype)


# ── Helpers de baixo nível ────────────────────────────────────────────────────

def _fmt_run(run, bold=False, italic=False, underline=False, size=None, color=None):
    run.font.name = FONTE
    run.font.size = size or TAMANHO
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.color.rgb = color or COR


def _normalizar_uf(valor):
    if valor is None:
        return None
    texto = str(valor).strip().upper()
    if len(texto) == 2 and texto.isalpha():
        return texto
    return UF_NOMES.get(texto)


def _signatarios_para_contexto(context):
    """Seleciona OAB apenas de UF declarada do processo originário."""
    uf_bruta = context.get('uf_processo_originario')
    if uf_bruta is None:
        uf_bruta = context.get('estado_processo_originario')
    if uf_bruta is None:
        return list(SIGNATARIOS)
    uf = _normalizar_uf(uf_bruta)
    if uf not in OAB_WANDERLEY_POR_UF:
        raise ValueError(
            f"UF do processo originário sem cadastro de OAB para Wanderley: {uf_bruta!r}. "
            "Informe uma UF cadastrada ou mantenha o padrão MG."
        )
    selecionados = list(SIGNATARIOS)
    nome, _, email, obs = selecionados[0]
    selecionados[0] = (nome, OAB_WANDERLEY_POR_UF[uf], email, obs)
    nome, _, email, obs = selecionados[1]
    selecionados[1] = (nome, OAB_FLAVIA_POR_UF.get(uf, OAB_FLAVIA_POR_UF['MG']), email, obs)
    return selecionados


def _adicionar_hyperlink(paragraph, texto, url):
    """Insere hyperlink externo azul e sublinhado, compatível com Word."""
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rid)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), FONTE)
    rfonts.set(qn('w:hAnsi'), FONTE)
    rpr.append(rfonts)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rpr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.append(underline)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rpr.append(sz)
    run.append(rpr)
    t = OxmlElement('w:t')
    t.text = texto
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _base_pf(pf, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
             left=Emu(0), first_line=Emu(0), hanging=None):
    pf.alignment = align
    pf.line_spacing_rule = line_rule
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = left
    if hanging is not None:
        # python-docx expressa "hanging indent" como first_line_indent negativo
        pf.first_line_indent = Emu(-hanging)
        pf.left_indent = left
    else:
        pf.first_line_indent = first_line


def _add_tab_stop(paragraph, position, alignment=WD_TAB_ALIGNMENT.LEFT):
    paragraph.paragraph_format.tab_stops.add_tab_stop(position, alignment)


def _add_bottom_border(paragraph, sz=4, color='000000', space=1):
    """Borda inferior no parágrafo inteiro (usada nos títulos I./II./III.),
    nunca sublinhado manual no texto. sz=4 (0,5pt) — conferido contra a
    Apelação revisada e os Embargos, as duas peças com numeração automática
    real (Correções.md, item 7: Agravo/Especificação usavam 6=0,75pt, que
    era o padrão manual divergente)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_top_border(paragraph, sz=6, color='000000', space=1):
    """Borda superior no parágrafo (usada na 1ª linha do rodapé — Apontamentos
    de melhoria 2026-07: rodapé estava sem linha separando do corpo)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(sz))
    top.set(qn('w:space'), str(space))
    top.set(qn('w:color'), color)
    pBdr.append(top)
    pPr.append(pBdr)


def _blank(doc, n=1, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE, keep_with_next=False):
    """Correções.md, item 5: 1,5 é o padrão para o separador entre blocos —
    entrelinha simples fica restrita às categorias com regra própria
    (título, citação, rodapé, assinaturas, endereçamento), que já formatam
    seus próprios parágrafos fora desta função."""
    for _ in range(n):
        p = doc.add_paragraph()
        _base_pf(p.paragraph_format, line_rule=line_rule)
        if keep_with_next:
            p.paragraph_format.keep_with_next = True


def _add_full_border(paragraph, sz=4, color='000000', space_tb=1, space_lr=4):
    """Borda nos 4 lados do parágrafo — replica o estilo 'ID Processo' real do
    escritório (styles.xml: top/bottom space=1, left/right space=4, sz=4).
    Parágrafos consecutivos com a mesma borda se fundem visualmente numa
    única caixa no Word (não desenha a borda interna entre eles)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side, space in (('top', space_tb), ('left', space_lr), ('bottom', space_tb), ('right', space_lr)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), str(space))
        el.set(qn('w:color'), color)
        pBdr.append(el)
    pPr.append(pBdr)


# ── Blocos do corpo (item 1 a 6 do plano) ─────────────────────────────────────

def bloco_abertura(doc, nome_parte, resto, nome_peca=None, resto_depois=''):
    """Item 1 — nome da parte em negrito + sublinhado SÓ neste parágrafo (a
    qualificação inicial) — única exceção à vedação geral de sublinhado do
    Manual RDAA §2.9 (decisão de 2026-08: praxe forense de identificar a
    parte na abertura exige negrito+sublinhado; o resto do documento
    continua proibido de usar sublinhado). Recuo de primeira linha em 2 cm
    (não é numerado).

    `nome_peca` é opcional: quando presente (ex.: "CONTRARRAZÕES AOS EMBARGOS
    DE DECLARAÇÃO"), sai em CAIXA ALTA + negrito logo após "apresentar" /
    "interpor" / etc., no meio do texto `resto` — passe a parte anterior do
    texto em `resto` e o restante (depois do nome da peça) em `resto_depois`."""
    p = doc.add_paragraph()
    r1 = p.add_run(nome_parte)
    _fmt_run(r1, bold=True, underline=True)
    r2 = p.add_run(resto)
    _fmt_run(r2)
    if nome_peca:
        r3 = p.add_run(nome_peca.upper())
        _fmt_run(r3, bold=True)
        r4 = p.add_run(resto_depois)
        _fmt_run(r4)
    _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE, first_line=CM2)
    return p


def bloco_titulo(doc, texto, gerenciador, sequencia='default', reiniciar=False):
    """Item 2 — título em maiúsculas e negrito, borda inferior, justificado,
    espaçamento 0/0, entrelinha simples. Numeração nativa (upperRoman + ".").

    Geometria (left=2cm, hanging=2cm, tabStop em 2cm) definida no nível
    de numeração (w:abstractNum/w:lvl/w:pPr) — o parágrafo não seta
    indent diretamente para não sobrescrever o nível."""
    p = doc.add_paragraph()
    p.style = doc.styles['RDAA Título 1']
    r = p.add_run(texto.upper())
    _fmt_run(r, bold=True)
    # Não chamar _base_pf: indent vem do nível de numeração
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _add_bottom_border(p)
    # Correções.md, "Título de tópico nunca pode ficar sozinho no final de
    # uma página": keepNext prende o título ao separador seguinte, que por
    # sua vez (ver construir_peca) também leva keep_with_next=True — a
    # cadeia se estende até o primeiro parágrafo real do tópico.
    pf.keep_with_next = True
    pf.keep_together = True
    num_id = gerenciador.num_id_para('titulo', sequencia, reiniciar)
    _aplicar_numeracao(p, num_id)
    return p


def bloco_numerado(doc, texto, gerenciador, sequencia='default', reiniciar=False):
    """Item 3 — parágrafo numerado (decimal + "."), tab fixo em 2 cm.
    Numeração nativa — geometria (left=2cm, hanging=2cm, tabStop em 2cm)
    definida no nível de numeração. Entrelinha 1,5, espaçamento 0/0.
    Suporta inline markup (<b>negrito</b>, <i>itálico</i>)."""
    p = doc.add_paragraph()
    p.style = doc.styles['RDAA Numerado']
    _adicionar_texto_formatado(p, texto)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    num_id = gerenciador.num_id_para('numerado', sequencia, reiniciar)
    _aplicar_numeracao(p, num_id)
    return p


def bloco_titulo2(doc, texto, gerenciador, sequencia='default', reiniciar=False):
    """Título Nível 2 (Subtópico) — Decimal (1., 2.), centralizado, negrito,
    só 1ª letra maiúscula (Manual §3.1 — não força .upper()), sem borda inferior,
    entrelinha 1,5, espaçamento 0/0, keep_with_next=True."""
    p = doc.add_paragraph()
    p.style = doc.styles['RDAA Título 2']
    _adicionar_texto_formatado(p, texto, bold_default=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.keep_with_next = True
    pf.keep_together = True
    num_id = gerenciador.num_id_para('titulo2', sequencia, reiniciar)
    _aplicar_numeracao(p, num_id)
    return p


def bloco_titulo3(doc, texto, gerenciador, sequencia='default', reiniciar=False):
    """Título Nível 3 (Sub-subtópico) — lowerLetter + ponto (a., b.), CAIXA ALTA,
    negrito, borda inferior, recuo 4cm, entrelinha 1,5, espaçamento 0/0, keep_with_next=True."""
    p = doc.add_paragraph()
    p.style = doc.styles['RDAA Título 3']
    r = p.add_run(texto.upper())
    _fmt_run(r, bold=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _add_bottom_border(p)
    pf.keep_with_next = True
    pf.keep_together = True
    num_id = gerenciador.num_id_para('titulo3', sequencia, reiniciar)
    _aplicar_numeracao(p, num_id)
    return p


def bloco_alinea(doc, texto, gerenciador, sequencia='default', reiniciar=False, nivel=0):
    """Item 4 — alínea multinível.
    nivel=0: lowerLetter + ")", marcador em 2 cm, texto em 3 cm.
    nivel=1: lowerRoman + ")", marcador em 3 cm, texto em 4 cm.
    Numeração nativa. Entrelinha 1,5, espaçamento 0/0."""
    p = doc.add_paragraph()
    p.style = doc.styles['RDAA Alínea']
    r = p.add_run(texto)
    _fmt_run(r)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    num_id = gerenciador.num_id_para('alinea', sequencia, reiniciar)
    _aplicar_numeracao(p, num_id, ilvl=nivel)
    return p


def bloco_documento(doc, texto, gerenciador, sequencia='default', reiniciar=False):
    """Bloco para lista de documentos anexos ("Doc. 01 - ...").
    Numeração nativa decimal pré-fixada com 'Doc. ' e hífen. Entrelinha 1.5."""
    p = doc.add_paragraph()
    p.style = doc.styles['RDAA Documento']
    r = p.add_run(texto)
    _fmt_run(r)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    num_id = gerenciador.num_id_para('documento', sequencia, reiniciar)
    _aplicar_numeracao(p, num_id)
    return p


def bloco_citacao(doc, texto, italic=False, bold=False):
    """Item 5 — recuo esquerdo 2 cm, sem recuo especial, justificado,
    entrelinha SIMPLES (não 1,5), espaçamento 0/0, sem numeração própria.
    Itálico/negrito são opcionais por bloco (nunca fixos) — a peça de origem
    transcrita não usava itálico. `texto` pode ter múltiplas linhas (\\n);
    cada linha vira um parágrafo próprio, sem parágrafo vazio entre elas."""
    paragrafos = []
    for linha in texto.split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['RDAA Citação']
        r = p.add_run(linha)
        _fmt_run(r, bold=bold, italic=italic, size=Pt(9))
        _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE, left=CM2, first_line=Emu(0))
        paragrafos.append(p)
    return paragrafos


def bloco_paragrafo_recuo(doc, texto, bold=False):
    """Parágrafos sem numeração explícita que ainda precisam do recuo de
    primeira linha em 2 cm (ex.: 'Nestes termos, aguarda deferimento.', a
    linha de local/data, o boilerplate de publicações)."""
    p = doc.add_paragraph()
    r = p.add_run(texto)
    _fmt_run(r, bold=bold)
    _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE, first_line=CM2)
    return p


# Item 6 (regra central): todos os blocos acima que precisam de tabulação usam
# _add_tab_stop com uma posição FIXA (CM2/CM3) — nunca "\t" solto. Sem isso o
# Word cai no intervalo padrão (~1,25 cm) e qualquer numeração de dois
# dígitos ou título mais largo perde o alinhamento vertical.


def _append_hidden_text(paragraph, text):
    """Anexar texto pesquisável sem alterar a aparência do parágrafo."""
    if not text:
        return
    run = paragraph.add_run(str(text))
    rpr = run._r.get_or_add_rPr()
    vanish = OxmlElement('w:vanish')
    rpr.append(vanish)


def bloco_figura(doc, image_path, legenda=None, width_cm=14.0, funcao_visual=None, texto_pesquisavel=None):
    """Insere uma imagem centralizada no documento com legenda opcional em 9pt/itálico.
    Nunca amplia a imagem além de sua dimensão física original."""
    if not os.path.exists(image_path):
        p = doc.add_paragraph()
        r = p.add_run(f"[Imagem não encontrada {image_path}]")
        _fmt_run(r, italic=True)
        _base_pf(p.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER)
        _append_hidden_text(p, _visual_search_text(funcao_visual, texto_pesquisavel))
        return p

    p = doc.add_paragraph()
    _base_pf(p.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))
    r = p.add_run()

    req_emu = Emu(round(width_cm * 360000))
    try:
        from PIL import Image
        with Image.open(image_path) as im:
            px_w, _ = im.size
            nat_emu = px_w * 9525  # 96 DPI: 1px = 9525 EMU
            final_emu = min(nat_emu, req_emu)
    except Exception:
        final_emu = req_emu

    r.add_picture(image_path, width=final_emu)

    paragrafos = [p]

    if legenda:
        p_leg = doc.add_paragraph()
        r_leg = p_leg.add_run(f"Figura {legenda}")
        _fmt_run(r_leg, italic=True, size=Pt(9))
        _base_pf(p_leg.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))
        paragrafos.append(p_leg)

    target = paragrafos[-1]
    _append_hidden_text(target, _visual_search_text(funcao_visual, texto_pesquisavel))
    return paragrafos


def _visual_search_text(funcao_visual=None, texto_pesquisavel=None, visual_tipo=None):
    parts = []
    if visual_tipo:
        parts.append(f"Visual Law {visual_tipo}")
    if funcao_visual:
        parts.append(f"Função visual {funcao_visual}")
    if texto_pesquisavel:
        parts.append(str(texto_pesquisavel))
    return " | ".join(parts)


def bloco_tabela(doc, cabecalho=None, linhas=None, alinhamentos=None):
    """Insere uma tabela genérica de dados no corpo do documento:
    - Centralizada como objeto no documento (WD_TABLE_ALIGNMENT.CENTER).
    - `cabecalho`: lista de strings para a primeira linha (em negrito).
    - `linhas`: lista de listas de strings para as demais linhas.
    - `alinhamentos`: lista opcional de alinhamentos por coluna ('left', 'center', 'right')."""
    if cabecalho is None:
        cabecalho = []
    if linhas is None:
        linhas = []

    cols_count = max(len(cabecalho), max((len(l) for l in linhas), default=0))
    if cols_count == 0:
        return None

    rows_count = (1 if cabecalho else 0) + len(linhas)
    tbl = doc.add_table(rows=rows_count, cols=cols_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    ALIGN_MAP = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }

    row_offset = 0
    if cabecalho:
        for col_idx, text in enumerate(cabecalho):
            cell = tbl.rows[0].cells[col_idx]
            p = cell.paragraphs[0]
            align = ALIGN_MAP.get(alinhamentos[col_idx], WD_ALIGN_PARAGRAPH.LEFT) if (alinhamentos and col_idx < len(alinhamentos)) else WD_ALIGN_PARAGRAPH.LEFT
            _adicionar_texto_formatado(p, text, bold_default=True)
            _base_pf(p.paragraph_format, align=align, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))
        row_offset = 1

    for row_idx, row_data in enumerate(linhas):
        for col_idx, text in enumerate(row_data):
            cell = tbl.rows[row_idx + row_offset].cells[col_idx]
            p = cell.paragraphs[0]
            align = ALIGN_MAP.get(alinhamentos[col_idx], WD_ALIGN_PARAGRAPH.LEFT) if (alinhamentos and col_idx < len(alinhamentos)) else WD_ALIGN_PARAGRAPH.LEFT
            _adicionar_texto_formatado(p, text)
            _base_pf(p.paragraph_format, align=align, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))

    return tbl


def bloco_visual(doc, visual_tipo, funcao_visual, texto_pesquisavel, cabecalho=None, linhas=None, alinhamentos=None):
    """Renderizar Visual Law tipado como tabela pesquisável."""
    tbl = bloco_tabela(doc, cabecalho, linhas, alinhamentos)
    if tbl is None:
        return None
    target = tbl.rows[0].cells[0].paragraphs[0]
    _append_hidden_text(target, _visual_search_text(funcao_visual, texto_pesquisavel, visual_tipo))
    return tbl


def bloco_sumula(doc, texto, italic=True):
    """Súmula da peça / síntese executiva — recuo esquerdo 2,5cm (1417 twips),
    recuo direito 2cm (1134 twips), itálico, justificado, entrelinha 1,5,
    sem borda, sem bullets, sem numeração."""
    p = doc.add_paragraph()
    _adicionar_texto_formatado(p, texto, italic_default=italic)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Emu(round(2.5 * 360000))
    pf.right_indent = Emu(round(2.0 * 360000))
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def bloco_assinaturas(doc, context=None):
    """Bloco explícito para inserção da tabela de assinaturas."""
    if not isinstance(context, dict):
        context = getattr(doc, '_rdaa_context', {})
    _inserir_tabela_assinaturas(doc, context=context)


def bloco_quadro_processual(doc, numero_processo=None, partes=None):
    """Insere o quadro de processo/partes em qualquer ponto da peça."""
    caixa_paragrafos = []
    if numero_processo:
        p = doc.add_paragraph()
        r_label = p.add_run('Processo ')
        _fmt_run(r_label, bold=True)
        r_num = p.add_run(numero_processo)
        _fmt_run(r_num, bold=False)
        _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))
        caixa_paragrafos.append(p)

    if partes:
        for linha in partes.split('\n'):
            p = doc.add_paragraph()
            if ':' in linha:
                rotulo, _, valor = linha.partition(':')
                r_label = p.add_run(rotulo + ' ')
                _fmt_run(r_label, bold=True)
                r_val = p.add_run(valor.lstrip())
                _fmt_run(r_val, bold=False)
            else:
                r = p.add_run(linha)
                _fmt_run(r, bold=False)
            _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))
            caixa_paragrafos.append(p)

    for p in caixa_paragrafos:
        _add_full_border(p)
    if caixa_paragrafos:
        _blank(doc, 2)
    return caixa_paragrafos


def bloco_inicio_razoes(doc, bloco, gerenciador):
    """Insere a transição para as Razões Recursais em recursos compostos:
    - Quebra de página.
    - Reinicia as sequências de numeração (títulos, numerados, alíneas).
    - Opcionalmente insere novo endereçamento e título de razões."""
    target = doc.add_page_break()
    seq = bloco.get('sequencia', 'razoes')
    gerenciador.num_id_para('titulo', sequencia=seq, reiniciar=True)
    gerenciador.num_id_para('numerado', sequencia=seq, reiniciar=True)
    gerenciador.num_id_para('alinea', sequencia=seq, reiniciar=True)

    if bloco.get('enderecamento'):
        p = doc.add_paragraph()
        target = p
        r = p.add_run(bloco['enderecamento'])
        _fmt_run(r, bold=True)
        _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE)
        _blank(doc, 2)

    if bloco.get('titulo_razoes'):
        p = doc.add_paragraph()
        p.style = doc.styles['RDAA Título Razões']
        target = p
        r = p.add_run(bloco['titulo_razoes'].upper())
        _fmt_run(r, bold=True)
        _base_pf(p.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE)
        _blank(doc, 2)
    return target


BLOCO_HANDLERS = {
    'abertura':           lambda doc, b, g: bloco_abertura(doc, b['nome_parte'], b['resto'],
                                                            b.get('nome_peca'), b.get('resto_depois', '')),
    'titulo':             lambda doc, b, g: bloco_titulo(doc, b['texto'], g,
                                                          b.get('sequencia', 'default'),
                                                          b.get('reiniciar', False)),
    'titulo2':            lambda doc, b, g: bloco_titulo2(doc, b['texto'], g,
                                                           b.get('sequencia', 'default'),
                                                           b.get('reiniciar', False)),
    'titulo3':            lambda doc, b, g: bloco_titulo3(doc, b['texto'], g,
                                                           b.get('sequencia', 'default'),
                                                           b.get('reiniciar', False)),
    'numerado':           lambda doc, b, g: bloco_numerado(doc, b['texto'], g,
                                                            b.get('sequencia', 'default'),
                                                            b.get('reiniciar', False)),
    'alinea':             lambda doc, b, g: bloco_alinea(doc, b['texto'], g,
                                                          b.get('sequencia', 'default'),
                                                          b.get('reiniciar', False),
                                                          b.get('nivel', 0)),
    'documento':          lambda doc, b, g: bloco_documento(doc, b['texto'], g,
                                                             b.get('sequencia', 'default'),
                                                             b.get('reiniciar', False)),
    'citacao':            lambda doc, b, g: bloco_citacao(doc, b['texto'], b.get('italic', False),
                                                           b.get('bold', False)),
    'paragrafo':          lambda doc, b, g: bloco_paragrafo_recuo(doc, b['texto'], b.get('bold', False)),
    'figura':             lambda doc, b, g: bloco_figura(
        doc, b['image_path'], b.get('legenda'), b.get('width_cm', 14.0),
        b.get('funcao_visual'), b.get('texto_pesquisavel')),
    'decisao_anotada':    lambda doc, b, g: bloco_figura(
        doc, b['image_path'], b.get('legenda'), b.get('width_cm', 14.0),
        b.get('funcao_visual', 'Destacar trecho explicitamente indicado'), b.get('texto_pesquisavel')),
    'tabela':             lambda doc, b, g: bloco_tabela(doc, b.get('cabecalho'), b.get('linhas'), b.get('alinhamentos')),
    'visual':             lambda doc, b, g: bloco_visual(
        doc, b['visual_tipo'], b['funcao_visual'], b['texto_pesquisavel'],
        b.get('cabecalho'), b.get('linhas'), b.get('alinhamentos')),
    'sumula':             lambda doc, b, g: bloco_sumula(doc, b['texto'], b.get('italic', True)),
    'inicio_razoes':      lambda doc, b, g: bloco_inicio_razoes(doc, b, g),
    'assinaturas':        lambda doc, b, g: bloco_assinaturas(doc),
    'quadro_processual': lambda doc, b, g: bloco_quadro_processual(doc, b.get('numero_processo'), b.get('partes')),
}

# Blocos que recebem parágrafo vazio depois (item 3/4: "Enter duplo", nunca
# spacing.after). A abertura também leva uma linha em branco antes do
# primeiro título, conforme o modelo real; "paragrafo" (fecho, publicações,
# data/local) é tratado à parte na função construir_peca, então não entra aqui.
BLOCOS_COM_BLANK_DEPOIS = {'abertura', 'titulo', 'titulo2', 'titulo3', 'numerado', 'alinea', 'documento', 'citacao', 'sumula', 'figura', 'decisao_anotada', 'tabela', 'visual'}
VISUAL_TIPOS = {'timeline', 'matrix', 'flow', 'confrontation'}


# Validação de entrada: erros de schema devem ser localizados antes de criar
# o DOCX. Isso evita KeyError/TypeError no meio da geração e não interfere em
# nenhum texto que já esteja válido segundo schema_blocos.md.
BLOCO_REQUIRED_FIELDS = {
    'abertura': ('nome_parte', 'resto'),
    'titulo': ('texto',),
    'titulo2': ('texto',),
    'titulo3': ('texto',),
    'numerado': ('texto',),
    'alinea': ('texto',),
    'documento': ('texto',),
    'citacao': ('texto',),
    'paragrafo': ('texto',),
    'figura': ('image_path',),
    'visual': ('visual_tipo', 'funcao_visual', 'texto_pesquisavel', 'linhas'),
    'decisao_anotada': ('image_path', 'texto_pesquisavel'),
    'sumula': ('texto',),
    'inicio_razoes': (),
    'assinaturas': (),
    'quadro_processual': (),
    'tabela': ('linhas',),
}


def _validar_texto_titulo(texto, local):
    texto = str(texto or '')
    ocorrencias = []
    for char, nome in TITULO_PROIBIDO_CHARS.items():
        if char in texto:
            ocorrencias.append(f'{nome} ({char})')
    if ocorrencias:
        raise ValueError(
            f'{local}: título contém caractere proibido: {", ".join(ocorrencias)}. '
            'Reescreva o título sem travessão e sem dois pontos.'
        )


def validar_contexto(context):
    """Valida o contrato mínimo do JSON antes de iniciar a geração."""
    if not isinstance(context, dict):
        raise ValueError('Contexto inválido: o JSON raiz deve ser um objeto.')
    _signatarios_para_contexto(context)
    blocos = context.get('blocos', [])
    if not isinstance(blocos, list):
        raise ValueError("Contexto inválido: 'blocos' deve ser uma lista.")
    for idx, bloco in enumerate(blocos, start=1):
        if not isinstance(bloco, dict):
            raise ValueError(f'Bloco {idx}: esperado objeto JSON, recebido {type(bloco).__name__}.')
        tipo = bloco.get('tipo')
        if tipo not in BLOCO_HANDLERS:
            raise ValueError(f"Bloco {idx}: tipo desconhecido {tipo!r}. Consulte references/schema_blocos.md.")
        for campo in BLOCO_REQUIRED_FIELDS.get(tipo, ()):
            valor = bloco.get(campo)
            if valor is None or (isinstance(valor, str) and not valor.strip()):
                raise ValueError(f"Bloco {idx} ({tipo}): campo obrigatório ausente ou vazio: '{campo}'.")
        if tipo == 'alinea' and bloco.get('nivel', 0) not in (0, 1):
            raise ValueError(f"Bloco {idx} (alinea): 'nivel' deve ser 0 ou 1.")
        if tipo in ('tabela', 'visual') and not isinstance(bloco.get('linhas'), list):
            raise ValueError(f"Bloco {idx} ({tipo}): 'linhas' deve ser uma lista.")
        if tipo == 'visual' and bloco.get('visual_tipo') not in VISUAL_TIPOS:
            raise ValueError(
                f"Bloco {idx} (visual): 'visual_tipo' deve ser um de {sorted(VISUAL_TIPOS)}."
            )
        if tipo in ('titulo', 'titulo2', 'titulo3'):
            _validar_texto_titulo(bloco.get('texto'), f'Bloco {idx} ({tipo})')
        if tipo == 'inicio_razoes' and bloco.get('titulo_razoes'):
            _validar_texto_titulo(bloco.get('titulo_razoes'), f'Bloco {idx} (titulo_razoes)')


# ── Cabeçalho (logo) ──────────────────────────────────────────────────────────

LOGO_PATH_PADRAO = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                                 'assets', 'Logo.jpg')


def _configurar_cabecalho(section, logo_path=None):
    """Insere a logo do escritório no cabeçalho, alinhada à direita, largura
    útil da página (17 cm — mesma largura usada na tabela de assinaturas).
    A imagem já traz embutida a régua horizontal inferior (Logo.jpg oficial),
    então não é preciso desenhar borda separada."""
    logo_path = logo_path or LOGO_PATH_PADRAO
    if not logo_path or not os.path.isfile(logo_path):
        # Correções.md, item 4: cabeçalho é obrigatório em toda peça RDAA —
        # gerar sem logo não é um documento válido, é uma falha silenciosa.
        raise FileNotFoundError(
            f"Cabeçalho RDAA obrigatório: logo não encontrada em {logo_path!r}."
        )
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    p = header.add_paragraph()
    run = p.add_run()
    run.add_picture(logo_path, width=TW(9639))  # 17 cm — largura útil (21 - 2 - 2 cm de margem)
    _base_pf(p.paragraph_format, align=WD_ALIGN_PARAGRAPH.RIGHT, line_rule=WD_LINE_SPACING.SINGLE)

    # Decisao 2026-08: o respiro entre a logo e o inicio do texto deve viver
    # DENTRO do cabecalho (paragrafo vazio aqui), nao como paragrafo em
    # branco no corpo antes do enderecamento — o corpo nao deve ter nenhum
    # enter extra antes da primeira linha de texto.
    p_respiro = header.add_paragraph()
    _base_pf(p_respiro.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE)


# ── Rodapé ────────────────────────────────────────────────────────────────────

def _add_page_number_field(para, field_code):
    # Correções.md, item 2: todos os runs do campo — inclusive os invisíveis
    # de fldChar/instrText — devem carregar Tahoma 8, não só o resultado.
    run = para.add_run(); _fmt_run(run, size=Pt(8))
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); run._r.append(fc)
    run2 = para.add_run(); _fmt_run(run2, size=Pt(8))
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = f' {field_code} '
    run2._r.append(it)
    run3 = para.add_run(); _fmt_run(run3, size=Pt(8))
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate'); run3._r.append(fc2)
    run4 = para.add_run('1'); _fmt_run(run4, size=Pt(8))
    run5 = para.add_run(); _fmt_run(run5, size=Pt(8))
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end'); run5._r.append(fc3)


def _configurar_rodape(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    # Item G1-2-7 (Apontamentos de melhoria 2026-07): rodapé sem linha
    # separando do corpo, e paginação centralizada — corrigido para linha
    # superior + paginação alinhada à direita.
    p1 = footer.add_paragraph()
    r1 = p1.add_run(RODAPE_L1); _fmt_run(r1, size=Pt(8))
    _base_pf(p1.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE)
    _add_top_border(p1)

    p2 = footer.add_paragraph()
    r2 = p2.add_run(RODAPE_L2); _fmt_run(r2, size=Pt(7), color=COR_DOURADA)
    _base_pf(p2.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE)

    p3 = footer.add_paragraph()
    r_pag = p3.add_run('Página '); _fmt_run(r_pag, size=Pt(8))
    _add_page_number_field(p3, 'PAGE')
    r_de = p3.add_run(' de '); _fmt_run(r_de, size=Pt(8))
    _add_page_number_field(p3, 'NUMPAGES')
    _base_pf(p3.paragraph_format, align=WD_ALIGN_PARAGRAPH.RIGHT, line_rule=WD_LINE_SPACING.SINGLE)


# ── Tabela de assinaturas ─────────────────────────────────────────────────────

def _remove_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    tb = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), 'auto')
        tb.append(b)
    tbl_pr.append(tb)


def _add_cell_margins(tbl, top=0, bottom=0, left=0, right=0):
    """Define w:tblCellMar (twips) — sem isso o Word usa a margem interna
    minima padrao e o texto de celulas adjacentes parece colado."""
    tbl_pr = tbl._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tbl_pr.append(mar)


def _inserir_tabela_assinaturas(doc, context=None):
    """Grade 2x2 sem bordas, conteúdo centralizado em cada célula (Correções.md,
    item 12: alinhamento centralizado é o padrão da Apelação revisada, adotada
    como referência visual). Sempre os 4 advogados, sempre nesta posição
    (Wanderley/Flávia na linha de cima, Alessandra/Ricardo Cesar embaixo)."""
    tbl = doc.add_table(rows=2, cols=2)
    _remove_borders(tbl)
    _add_cell_margins(tbl, top=57, bottom=170, left=113, right=113)  # ~0,3cm/0,5cm/0,2cm/0,2cm — sem isso o texto encosta na borda da celula ("assinaturas coladas")

    for idx, (nome, oab, email, obs) in enumerate(_signatarios_para_contexto(context or {})):
        row, col = divmod(idx, 2)
        cell = tbl.rows[row].cells[col]

        p_nome = cell.paragraphs[0]
        r = p_nome.add_run(nome)
        _fmt_run(r, bold=True)
        _base_pf(p_nome.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE)

        p_oab = cell.add_paragraph()
        r = p_oab.add_run(oab)
        _fmt_run(r)
        _base_pf(p_oab.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE)

        # E-mail segue o padrão visual fornecido por Ricardo, com hiperlink
        # externo azul e sublinhado nativo do Word.
        p_email = cell.add_paragraph()
        _adicionar_hyperlink(p_email, email, f'mailto:{email}')
        _base_pf(p_email.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE)

        if obs:
            p_obs = cell.add_paragraph()
            r = p_obs.add_run(f'({obs})')
            _fmt_run(r, italic=True)
            _base_pf(p_obs.paragraph_format, align=WD_ALIGN_PARAGRAPH.CENTER, line_rule=WD_LINE_SPACING.SINGLE)

    tbl.autofit = False
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Emu(round(8.5 * 360000))  # metade dos 17 cm úteis


# ── Documento principal ───────────────────────────────────────────────────────

def construir_peca(context: dict, output_path: str) -> str:
    validar_contexto(context)
    doc = Document()
    doc._rdaa_context = context

    sec = doc.sections[0]
    sec.page_height = Emu(round(29.7 * 360000))
    sec.page_width = Emu(round(21.0 * 360000))
    sec.top_margin = Emu(round(2 * 360000))
    sec.bottom_margin = Emu(round(2 * 360000))
    sec.left_margin = Emu(round(2 * 360000))
    sec.right_margin = Emu(round(2 * 360000))
    sec.header_distance = Emu(round(1 * 360000))
    sec.footer_distance = Emu(round(1 * 360000))
    sec.different_first_page_header_footer = False

    sn = doc.styles['Normal']
    sn.font.name = FONTE
    sn.font.size = TAMANHO
    sn.font.color.rgb = COR
    _base_pf(sn.paragraph_format, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE)

    _configurar_rodape(sec)
    _configurar_cabecalho(sec, context.get('logo_path'))

    # Estilos RDAA nomeados no Word (Fase 4)
    _criar_estilos_rdaa(doc)

    # Numeração nativa: registra listas base e gerenciador de sequências
    abstract_ids, ct_numbering = _preparar_numeracao(doc)
    gerenciador = _GerenciadorSequencias(ct_numbering, abstract_ids)
    gerenciador_notas = _GerenciadorNotasRodape(doc)

    # Decisao 2026-08: o respiro entre cabecalho e corpo agora vive dentro
    # do cabecalho (ver _configurar_cabecalho) — nao adicionar enter extra
    # aqui. Endereçamento é o primeiro paragrafo do corpo, sem linha em
    # branco antes. — espaçamento simples (não 1,5, exceção já prevista pelo
    # padrão RDAA) e 2 linhas em branco depois (estava com 1,5 e só 1 enter).
    if context.get('enderecamento'):
        p = doc.add_paragraph()
        r = p.add_run(context['enderecamento'])
        _fmt_run(r, bold=True)
        _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE)
        _blank(doc, 2)

    # Número do processo / partes — caixa com borda nos 4 lados (estilo real
    # "ID Processo"): rótulo em negrito, valor normal, uma linha por dado.
    # Parágrafos consecutivos com a mesma borda se fundem numa caixa só.
    caixa_paragrafos = []
    if context.get('numero_processo'):
        p = doc.add_paragraph()
        r_label = p.add_run('Processo ')
        _fmt_run(r_label, bold=True)
        r_num = p.add_run(context['numero_processo'])
        _fmt_run(r_num, bold=False)
        _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))
        caixa_paragrafos.append(p)

    # Cada parte em sua própria linha, sem parágrafo vazio entre elas — os
    # quatro .docx de referência confundem "um enter entre as partes" com
    # "cada parte em sua própria linha", não com linha vazia (Correções.md,
    # item 10). Ordem de polo ativo antes do passivo é responsabilidade de
    # quem monta `partes` (ver references/schema_blocos.md) — não é uma
    # regra desta função.
    if context.get('partes'):
        for linha in context['partes'].split('\n'):
            p = doc.add_paragraph()
            if ':' in linha:
                rotulo, _, valor = linha.partition(':')
                r_label = p.add_run(rotulo + ' ')
                _fmt_run(r_label, bold=True)
                r_val = p.add_run(valor.lstrip())
                _fmt_run(r_val, bold=False)
            else:
                r = p.add_run(linha)
                _fmt_run(r, bold=False)
            _base_pf(p.paragraph_format, line_rule=WD_LINE_SPACING.SINGLE, first_line=Emu(0))
            caixa_paragrafos.append(p)

    for p in caixa_paragrafos:
        _add_full_border(p)
    if caixa_paragrafos:
        _blank(doc, 2)

    # Corpo — lista de blocos tipados.
    # Se o bloco 'assinaturas' for o ULTIMO item da lista, ele nao entra no
    # loop aqui — sai reservado para depois do fecho fixo (publicacoes/
    # "aguarda deferimento"/data), que sao sempre anexados apos o loop.
    # Sem essa checagem, um 'assinaturas' colocado por ultimo no corpo fica
    # preso ANTES do fecho, porque o fecho e incondicional (bug real:
    # assinatura aparecia antes de "Nestes termos, aguarda deferimento.").
    # 'assinaturas' no MEIO da lista (recurso composto: interposicao +
    # razoes) continua funcionando normalmente, sem entrar nesse desvio.
    blocos_corpo = context.get('blocos', [])
    assinatura_final_explicita = bool(blocos_corpo) and blocos_corpo[-1].get('tipo') == 'assinaturas'
    if assinatura_final_explicita:
        blocos_corpo = blocos_corpo[:-1]

    semantic_occurrences = {}
    for bloco in blocos_corpo:
        tipo = bloco.get('tipo')
        handler = BLOCO_HANDLERS.get(tipo)
        if handler is None:
            raise ValueError(f"Tipo de bloco desconhecido: {tipo!r}")
        res = handler(doc, bloco, gerenciador)
        _mark_semantic_ids(res, bloco, semantic_occurrences)
        if bloco.get('nota_rodape'):
            p_alvo = res[-1] if isinstance(res, list) else res
            if p_alvo is not None:
                gerenciador_notas.adicionar_nota(p_alvo, bloco['nota_rodape'])
        if tipo in BLOCOS_COM_BLANK_DEPOIS:
            _blank(doc, keep_with_next=(tipo == 'titulo'))

    # Publicações (boilerplate opcional)
    if context.get('publicacoes', True):
        bloco_paragrafo_recuo(doc, context.get('publicacoes_texto', PUBLICACOES_PADRAO))
        _blank(doc)

    # Fecho
    bloco_paragrafo_recuo(doc, context.get('fecho', 'Nestes termos, aguarda deferimento.'))
    _blank(doc)

    # Data / local
    if context.get('data_local'):
        bloco_paragrafo_recuo(doc, context['data_local'])
        _blank(doc)

    # Assinaturas: sempre por ultimo, depois do fecho — automatica se nao
    # houver bloco explicito, ou o bloco explicito que foi retirado do loop
    # acima por ser o ultimo item do corpo.
    tem_assinaturas_no_meio = any(b.get('tipo') == 'assinaturas' for b in blocos_corpo)
    if assinatura_final_explicita or (context.get('assinatura_automatica_final', True) and not tem_assinaturas_no_meio):
        _inserir_tabela_assinaturas(doc, context=context)

        # Finaliza e grava as notas de rodapé reais no pacote OPC
    gerenciador_notas.finalizar()

    output_path = os.path.abspath(output_path)
    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)
    # Salva em arquivo temporário no mesmo diretório e troca de forma atômica
    # (os.replace) — se o processo for interrompido no meio da gravação, um
    # candidato .docx já existente em output_path não fica truncado/corrompido.
    fd, tmp_path = tempfile.mkstemp(prefix=f".{os.path.basename(output_path)}.", suffix=".tmp", dir=output_dir)
    os.close(fd)
    try:
        doc.save(tmp_path)
        os.replace(tmp_path, output_path)
    except Exception:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise
    return output_path


def main():
    parser = argparse.ArgumentParser(description='Gera .docx RDAA nativo (python-docx puro)')
    parser.add_argument('--context', default='/tmp/rdaa_context.json')
    parser.add_argument('--output', default='outputs/peca_final.docx')
    args = parser.parse_args()

    with open(args.context, encoding='utf-8') as f:
        data = json.load(f)

    output_path = construir_peca(data, args.output)
    print(f'Gerado: {output_path}')
    return output_path


if __name__ == '__main__':
    main()
