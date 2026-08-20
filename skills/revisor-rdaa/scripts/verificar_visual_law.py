#!/usr/bin/env python3
"""Validação objetiva de elementos Visual Law no contexto e no DOCX."""

from __future__ import annotations

import argparse
import hashlib
import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise SystemExit(f"python-docx necessário: {exc}")

VISUAL_TIPOS = {"timeline", "matrix", "flow", "confrontation"}


def _docx_text(document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def verify_visual_law(docx_path: Path | str, context: dict[str, Any]) -> dict[str, Any]:
    findings: list[dict[str, Any]] = []
    blocks = [
        (index, block)
        for index, block in enumerate(context.get("blocos", []), start=1)
        if isinstance(block, dict) and block.get("tipo") in {"visual", "figura", "decisao_anotada"}
    ]
    for index, block in blocks:
        tipo = block.get("tipo")
        if tipo == "visual":
            visual_tipo = block.get("visual_tipo")
            if visual_tipo not in VISUAL_TIPOS:
                findings.append({"kind": "visual_type_invalid", "severity": "erro", "message": f"Bloco {index}: tipo Visual Law inválido"})
            for field in ("funcao_visual", "texto_pesquisavel"):
                if not str(block.get(field) or "").strip():
                    findings.append({"kind": "visual_field_missing", "severity": "erro", "message": f"Bloco {index}: campo obrigatório ausente: {field}"})
            if not isinstance(block.get("linhas"), list) or not block.get("linhas"):
                findings.append({"kind": "visual_rows_missing", "severity": "erro", "message": f"Bloco {index}: Visual Law deve possuir linhas"})
        elif block.get("texto_pesquisavel") is not None and not str(block.get("texto_pesquisavel")).strip():
            findings.append({"kind": "figure_search_text_empty", "severity": "alerta", "message": f"Bloco {index}: texto pesquisável vazio"})
        if tipo == "decisao_anotada":
            manifest_path = Path(str(block.get("annotation_manifest") or ""))
            if block.get("annotation_manifest") and not manifest_path.is_file():
                findings.append({"kind": "annotation_manifest_missing", "severity": "erro", "message": f"Bloco {index}: manifesto de anotação não encontrado"})
            elif block.get("annotation_manifest"):
                try:
                    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                    image_path = Path(str(block.get("image_path") or ""))
                    if image_path.is_file() and manifest.get("output_sha256"):
                        digest = hashlib.sha256(image_path.read_bytes()).hexdigest()
                        if digest != manifest.get("output_sha256"):
                            findings.append({"kind": "annotation_output_hash_mismatch", "severity": "erro", "message": f"Bloco {index}: hash da imagem anotada difere do manifesto"})
                except (OSError, json.JSONDecodeError):
                    findings.append({"kind": "annotation_manifest_invalid", "severity": "erro", "message": f"Bloco {index}: manifesto de anotação inválido"})

    text = _docx_text(Document(str(docx_path))) if Path(docx_path).is_file() else ""
    for index, block in blocks:
        search_text = str(block.get("texto_pesquisavel") or "").strip()
        if search_text and search_text not in text:
            findings.append({
                "kind": "visual_search_text_missing",
                "severity": "erro",
                "message": f"Texto pesquisável do bloco {index} não localizado no DOCX",
                "localizacao": f"contexto.blocos[{index}]",
            })

    blocking = [finding for finding in findings if finding.get("severity") == "erro"]
    return {
        "status": "BLOCK" if blocking else "PASS",
        "docx": str(docx_path),
        "visual_blocks": len(blocks),
        "findings": findings,
        "counts": {
            "total": len(findings),
            "blocking": len(blocking),
            "alerts": sum(1 for finding in findings if finding.get("severity") == "alerta"),
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Verificar Visual Law rastreável no DOCX")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--context", type=Path, required=True)
    args = parser.parse_args()
    context = json.loads(args.context.read_text(encoding="utf-8"))
    report = verify_visual_law(args.docx, context)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
