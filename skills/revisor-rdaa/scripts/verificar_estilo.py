#!/usr/bin/env python3
"""
verificar_estilo.py — QA automatico de cadencia/estilo para pecas RDAA.

Espelha o verificar_formatacao.py do formatar-peca: em vez de confiar em
leitura estrutural de uma LLM para contar travessao, ponto-e-virgula e
tricolon de negacao (checklist-3-estilometria.md), este script CONTA.
Itens com regra objetiva entram no exit code. A peça final não pode conter
travessão (proibido sem exceção), ponto-e-vírgula fora de lista/alínea
(permitido apenas em parágrafos com estilo RDAA Numerado/RDAA Alínea, ex.:
pedidos em cascata), tricolon de negação, abertura defensiva recorrente,
dois-pontos ou aposto explicativo entre parênteses.
Marcadores isolados de lista
como (a), (i) e (1) são permitidos porque não são explicações.

Uso:
    python3 verificar_estilo.py caminho/para/peca.docx
    python3 verificar_estilo.py caminho/para/peca.txt

Sai com codigo 0 se nenhuma regra objetiva for violada. Sai com 1 caso
contrario. A lista de ocorrencias de dois-pontos continua sendo impressa para
facilitar o diagnóstico.
"""

import re
import sys

# ponytail: mesmo fix de construir_peca.py/verificar_formatacao.py/qa_gate.py
# — sem isso, mensagem com acento sai como mojibake quando outro script
# (qa_gate.py, ou um teste) captura este stdout no Windows.
for _stream in (sys.stdout, sys.stderr):
    if hasattr(_stream, "reconfigure"):
        _stream.reconfigure(encoding="utf-8")

JANELA_PAGINA = 12  # aprox. paragrafos por pagina (nao ha paginacao real fora do Word)


_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _tem_borda_completa(paragraph):
    """Detecta a caixa Processo/partes (borda nos 4 lados, ver
    _add_full_border em construir_peca.py) — ela usa dois-pontos por
    convenção própria (rótulo: valor) e não é prosa argumentativa."""
    pbdr = paragraph._p.find(f'.//{{{_W_NS}}}pPr/{{{_W_NS}}}pBdr')
    if pbdr is None:
        return False
    lados = {'top', 'left', 'bottom', 'right'}
    presentes = {child.tag.rsplit('}', 1)[-1] for child in pbdr}
    return lados.issubset(presentes)


def _paragraphs_from_docx(path):
    import docx
    doc = docx.Document(path)
    paragraphs = []
    estilos = []
    bordas = []
    seen_paragraphs = set()

    def append_paragraphs(items):
        for paragraph in items:
            marker = id(paragraph._p)
            if marker in seen_paragraphs:
                continue
            seen_paragraphs.add(marker)
            paragraphs.append(paragraph.text)
            estilos.append(paragraph.style.name if paragraph.style is not None else None)
            bordas.append(_tem_borda_completa(paragraph))

    def append_table(table):
        for row in table.rows:
            for cell in row.cells:
                append_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    append_table(nested_table)

    append_paragraphs(doc.paragraphs)
    for table in doc.tables:
        append_table(table)
    return paragraphs, estilos, bordas


def _paragraphs_from_txt(path):
    with open(path, encoding='utf-8') as f:
        linhas = [l.rstrip('\n') for l in f]
    return linhas, [None] * len(linhas), [False] * len(linhas)


def carregar_paragrafos(path):
    """Retorna (paragrafos, estilos, bordas). `estilos[i]` é o nome do
    estilo do parágrafo i no DOCX (ex.: 'RDAA Numerado', 'RDAA Alínea') ou
    None quando a fonte é .txt ou o parágrafo não tem estilo nomeado.
    `bordas[i]` indica se o parágrafo tem borda nos 4 lados (caixa
    Processo/partes)."""
    if path.lower().endswith('.docx'):
        return _paragraphs_from_docx(path)
    return _paragraphs_from_txt(path)


def _split_sentencas(paragrafo):
    # Split simples por . ! ? seguido de espaço/fim — nao trata abreviacoes,
    # suficiente para contar ocorrencias de ; dentro do mesmo periodo.
    return [s for s in re.split(r'(?<=[.!?])\s+', paragrafo) if s.strip()]


def checar_travessao(paragrafos):
    # Travessão é proibido na peça final, sem exceção (checklist-3, item J) —
    # não é um limite de recorrência, qualquer ocorrência bloqueia.
    problemas = []
    candidatos = []
    for i, p in enumerate(paragrafos):
        if '—' in p:
            candidatos.append((i, p[:100]))
            problemas.append(
                f"Paragrafo {i}: travessao proibido na peca final, sem excecao — "
                f"reescrever com virgula, ponto ou conectivo: {p.strip()[:100]!r}"
            )
    return problemas, candidatos


_ESTILOS_LISTA = {"rdaa numerado", "rdaa alínea"}


def checar_ponto_e_virgula(paragrafos, estilos=None):
    # Ponto-e-vírgula é permitido apenas em parágrafos de lista/alínea (ex.:
    # pedidos em cascata); em prosa corrida, qualquer ocorrência bloqueia.
    problemas = []
    if estilos is None:
        estilos = [None] * len(paragrafos)
    for i, p in enumerate(paragrafos):
        estilo = (estilos[i] or "").strip().casefold()
        if estilo in _ESTILOS_LISTA:
            continue
        if ';' in p:
            problemas.append(
                f"Paragrafo {i}: ponto-e-virgula fora de lista/alinea proibido — "
                f"reescrever com ponto ou conectivo: {p.strip()[:100]!r}"
            )
    return problemas


_ABERTURA_DEFENSIVA = re.compile(
    r'^\s*(?:não\s+se\s+(?:pretende|busca|trata|ignora|desconhece|cuida|discute)|'
    r'a\s+questão\s+não\s+está\s+em|o\s+que\s+não\s+se\s+(?:pretende|busca|trata))\b',
    re.IGNORECASE,
)


def _parece_titulo(paragrafo):
    letras = [char for char in paragrafo if char.isalpha()]
    if not letras:
        return False
    return len(paragrafo.split()) <= 12 and sum(char.isupper() for char in letras) / len(letras) > 0.9


def listar_aberturas_defensivas(paragrafos):
    ocorrencias = []
    for i, paragrafo in enumerate(paragrafos):
        if _parece_titulo(paragrafo):
            continue
        primeira_frase = re.split(r'(?<=[.!?])\s+', paragrafo.strip(), maxsplit=1)[0]
        match = _ABERTURA_DEFENSIVA.match(primeira_frase)
        if match:
            formula = re.sub(r'\s+', ' ', match.group(0).strip()).lower()
            ocorrencias.append((i, formula, primeira_frase[:140]))
    return ocorrencias


def checar_aberturas_defensivas(paragrafos):
    ocorrencias = listar_aberturas_defensivas(paragrafos)
    if len(ocorrencias) < 3:
        formulas = {}
        for _, formula, _ in ocorrencias:
            formulas[formula] = formulas.get(formula, 0) + 1
        if max(formulas.values(), default=0) < 2:
            return []
    indices = [item[0] for item in ocorrencias]
    formulas = {}
    for _, formula, _ in ocorrencias:
        formulas[formula] = formulas.get(formula, 0) + 1
    repetidas = [f'{formula} ({quantidade}x)' for formula, quantidade in formulas.items() if quantidade >= 2]
    detalhe = f' fórmulas repetidas {repetidas}.' if repetidas else '.'
    return [
        f'Abertura defensiva recorrente em {len(ocorrencias)} parágrafo(s), nos parágrafos {indices}.{detalhe} '
        'Reescrever positivamente ou justificar a função argumentativa indispensável.'
    ]


def checar_tricolon_negacao(paragrafos):
    problemas = []
    padrao = re.compile(
        r'\bn[aã]o\b[^.;:]*?,\s*n[aã]o\b[^.;:]*?(?:,|\se\s)\s*n[aã]o\b|'
        r'\bnem\b[^.;:]*?\bnem\b[^.;:]*?\bnem\b',
        re.IGNORECASE,
    )
    for i, p in enumerate(paragrafos):
        for m in padrao.finditer(p):
            problemas.append(
                f"Paragrafo {i}: tricolon/tetracolon de negacao — {m.group(0)[:100]!r}"
            )
    return problemas


def checar_dois_pontos(paragrafos, bordas=None, estilos=None):
    # A caixa Processo/partes usa "rótulo: valor" por convenção própria
    # (ex.: "Autora: Trivale..."), não é prosa argumentativa — ver
    # _tem_borda_completa. Citação literal (estilo "RDAA Citação") também é
    # isenta pelo mesmo motivo de checar_aposto_explicativo — ementa real de
    # tribunal frequentemente traz dois-pontos ("EMENTA:", "Tema 858:") e a
    # transcrição verbatim não pode reescrever isso.
    problemas = []
    if bordas is None:
        bordas = [False] * len(paragrafos)
    if estilos is None:
        estilos = [None] * len(paragrafos)
    for i, p in enumerate(paragrafos):
        if bordas[i]:
            continue
        if (estilos[i] or "").strip().casefold() == "rdaa citação":
            continue
        if ':' in p:
            problemas.append(
                f"Paragrafo {i}: dois-pontos proibido na peça final — reescrever com ponto, vírgula ou conectivo."
            )
    return problemas


_MARCADOR_PARENTESES = re.compile(r'^(?:[a-z]{1,3}|[ivxlcdm]{1,8}|\d{1,4})$', re.IGNORECASE)

# Marcadores institucionais fixos do gerador de peças (ex.: carimbo de
# assinatura eletrônica) — não são aposto explicativo em prosa argumentativa,
# são elementos estruturais exigidos por verificar_formatacao.py.
_MARCADORES_INSTITUCIONAIS_PARENTESES = {"assinado eletronicamente"}


def checar_aposto_explicativo(paragrafos, estilos=None):
    # Citação literal de jurisprudência (estilo "RDAA Citação") é transcrita
    # verbatim — se o tribunal de origem usou parênteses na própria ementa
    # (ex.: "(art. 966, VIII e § 1º, do CPC)"), a peça não pode reescrever
    # isso sem alterar a citação, o que a regra de redação proíbe. A regra de
    # aposto explicativo continua valendo para a prosa argumentativa própria.
    problemas = []
    if estilos is None:
        estilos = [None] * len(paragrafos)
    padrao = re.compile(r'\(([^()\r\n]*)\)')
    for i, p in enumerate(paragrafos):
        estilo = (estilos[i] or "").strip().casefold()
        if estilo == "rdaa citação":
            continue
        for match in padrao.finditer(p):
            conteudo = match.group(1).strip()
            if not conteudo or _MARCADOR_PARENTESES.fullmatch(conteudo):
                continue
            if conteudo.casefold() in _MARCADORES_INSTITUCIONAIS_PARENTESES:
                continue
            problemas.append(
                f"Paragrafo {i}: aposto explicativo entre parênteses proibido — reescrever em frase própria."
            )
        # Aposto entre travessões pareados já é coberto por checar_travessao,
        # que agora bloqueia qualquer travessão sem exceção.
    return problemas


def listar_dois_pontos(paragrafos, bordas=None, estilos=None):
    if bordas is None:
        bordas = [False] * len(paragrafos)
    if estilos is None:
        estilos = [None] * len(paragrafos)
    candidatos = []
    for i, p in enumerate(paragrafos):
        if bordas[i] or (estilos[i] or "").strip().casefold() == "rdaa citação":
            continue
        for s in _split_sentencas(p):
            if ':' in s:
                candidatos.append((i, s.strip()[:120]))
    return candidatos


_NUMERACAO_INICIAL = re.compile(r'^\s*(?:\d{1,4}|[ivxlcdm]{1,8})\s*[.)]\s*', re.IGNORECASE)

# Conectivos de transição não mudam a função da abertura — "Assim, X faz Y"
# e "X faz Y" abrem do mesmo jeito para efeito de variedade de redação.
_CONECTIVOS_TRANSICAO = ('assim', 'além disso', 'alem disso', 'nesse sentido', 'por sua vez')

# Demonstrativo + substantivo genérico ("esse contexto", "essa circunstância",
# "essa situação", "este cenário") tem a mesma função em qualquer combinação
# — normaliza pra um marcador comum antes de comparar.
_DEMONSTRATIVO_GENERICO = re.compile(
    r'^(?:esse|essa|este|esta|aquele|aquela)\s+'
    r'(?:contexto|circunst[aâ]ncia|situa[cç][aã]o|cen[aá]rio)\b',
    re.IGNORECASE,
)

_PALAVRAS_VAZIAS = {'de', 'da', 'do', 'das', 'dos', 'que', 'a', 'o', 'e'}


def _assinatura_abertura(paragrafo):
    """Normaliza a abertura de um parágrafo argumentativo só para efeito de
    comparação de QA — nunca altera o texto entregue. Remove numeração e
    conectivo de transição, reduz demonstrativo+substantivo genérico a um
    marcador comum, e reduz o resto ao sujeito/verbo da primeira oração útil
    (aproximado pelas primeiras palavras de conteúdo). Duas aberturas com a
    mesma função e estrutura (ex.: "Esse contexto demonstra que..." e "Essa
    circunstância demonstra que...") produzem a mesma assinatura mesmo com
    palavras diferentes — trocar sinônimo não basta para variar a abertura."""
    primeira_frase = re.split(r'(?<=[.!?])\s+', paragrafo.strip(), maxsplit=1)[0]
    texto = _NUMERACAO_INICIAL.sub('', primeira_frase).strip()

    for conectivo in _CONECTIVOS_TRANSICAO:
        padrao = re.compile(r'^' + re.escape(conectivo) + r'\s*,?\s*', re.IGNORECASE)
        substituido = padrao.sub('', texto)
        if substituido != texto:
            texto = substituido.strip()
            break

    texto = _DEMONSTRATIVO_GENERICO.sub('ISSO', texto)

    # Sujeito+verbo da primeira oração útil: corta no primeiro "que" ou
    # vírgula (fronteira comum entre sujeito/verbo e o complemento, que é a
    # parte que deveria variar de um parágrafo pro outro). Sem essa
    # fronteira, o conteúdo que vem depois — legitimamente diferente entre
    # parágrafos — poluiria a assinatura e mascararia a repetição real.
    corte = re.search(r'\bque\b|,', texto, flags=re.IGNORECASE)
    nucleo = texto[:corte.start()] if (corte and corte.start() > 0) else texto

    palavras = [w for w in re.findall(r"[\wÀ-ÿ]+", nucleo.lower()) if w not in _PALAVRAS_VAZIAS]
    return ' '.join(palavras[:6])


def checar_aberturas_consecutivas(paragrafos, minimo=2):
    # Parágrafos argumentativos consecutivos cuja abertura tem a mesma
    # função e estrutura (mesmo sujeito/verbo, mesmo conectivo, mesmo
    # demonstrativo genérico) leem mal, ainda que o resto da frase varie —
    # contagem objetiva sobre a assinatura normalizada, não sobre a palavra
    # literal. A comparação é só de QA: não normaliza nem reescreve a peça.
    problemas = []
    grupo = []
    assinatura_atual = None

    def fechar_grupo():
        if assinatura_atual and len(grupo) >= minimo:
            problemas.append(
                f"Paragrafos {grupo} tem aberturas estruturalmente equivalentes "
                f"(assinatura {assinatura_atual!r}) — variar sujeito/verbo/fato, nao so as palavras."
            )

    for i, p in enumerate(paragrafos):
        texto = p.strip()
        if not texto or _parece_titulo(texto):
            fechar_grupo()
            grupo, assinatura_atual = [], None
            continue
        assinatura = _assinatura_abertura(texto)
        if not assinatura:
            fechar_grupo()
            grupo, assinatura_atual = [], None
            continue
        if assinatura == assinatura_atual:
            grupo.append(i)
        else:
            fechar_grupo()
            grupo, assinatura_atual = [i], assinatura
    fechar_grupo()
    return problemas


def checar_aberturas_repetidas(paragrafos, minimo_palavras=4, limite=3):
    problemas = []
    aberturas = {}
    for i, p in enumerate(paragrafos):
        palavras = p.strip().split()
        if len(palavras) < minimo_palavras:
            continue
        chave = ' '.join(palavras[:minimo_palavras]).lower()
        aberturas.setdefault(chave, []).append(i)
    for chave, idxs in aberturas.items():
        if len(idxs) >= limite:
            problemas.append(
                f"Abertura repetida {limite}+ vezes ({chave!r}) nos paragrafos {idxs}."
            )
    return problemas


def checar(path):
    paragrafos, estilos, bordas = carregar_paragrafos(path)

    problemas = []
    trav_problemas, trav_candidatos = checar_travessao(paragrafos)
    problemas += trav_problemas
    problemas += checar_ponto_e_virgula(paragrafos, estilos)
    problemas += checar_tricolon_negacao(paragrafos)
    problemas += checar_aberturas_defensivas(paragrafos)
    problemas += checar_dois_pontos(paragrafos, bordas, estilos)
    problemas += checar_aposto_explicativo(paragrafos, estilos)
    problemas += checar_aberturas_repetidas(paragrafos)
    problemas += checar_aberturas_consecutivas(paragrafos)

    dois_pontos = listar_dois_pontos(paragrafos, bordas, estilos)

    return problemas, dois_pontos


def main():
    if len(sys.argv) != 2:
        print("Uso: python3 verificar_estilo.py caminho/para/peca.docx", file=sys.stderr)
        sys.exit(2)

    path = sys.argv[1]
    problemas, dois_pontos = checar(path)

    if problemas:
        print(f"[ERRO] {len(problemas)} problema(s) de cadencia/estilo em {path}:")
        for p in problemas:
            print(f"  - {p}")
    else:
        print(f"[OK] Nenhum item de contagem obrigatoria violado em {path}")

    aberturas_defensivas = listar_aberturas_defensivas(carregar_paragrafos(path)[0])
    if aberturas_defensivas:
        print(f"\n[INFO] {len(aberturas_defensivas)} abertura(s) defensiva(s) identificada(s) para revisão funcional:")
        for i, formula, trecho in aberturas_defensivas:
            print(f"  - par.{i}: {formula} — {trecho}")

    print(f"\n[INFO] {len(dois_pontos)} sentenca(s) com dois-pontos encontradas:")
    for i, s in dois_pontos[:200]:
        print(f"  - par.{i}: {s}")

    sys.exit(1 if problemas else 0)


if __name__ == '__main__':
    main()
