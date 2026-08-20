#!/usr/bin/env python3
"""Regressões da ligação entre contexto semântico e DOCX."""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / "skills" / "formatar-peca" / "scripts" / "construir_peca.py"
PUBLISH = ROOT / "skills" / "revisor-rdaa" / "scripts" / "publicar_docx.py"
FIXTURE = ROOT / "tests" / "fixtures" / "context_happy.json"
SCRIPT_DIR = ROOT / "skills" / "revisor-rdaa" / "scripts"
sys.path.insert(0, str(SCRIPT_DIR))

from verificar_semantica_docx import actual_ids_from_docx, verify_docx_semantics  # noqa: E402


def _context_with_ids() -> dict:
    context = json.loads(FIXTURE.read_text(encoding="utf-8"))
    context["numero_processo"] = "3333333-33.3333.8.26.3333"
    context["teses"] = [{"id": "T-1", "texto": "Tese explicitamente fornecida"}]
    context["blocos"][2].update(
        {
            "id": "B-1",
            "thesis_ids": ["T-1"],
            "semantic_ids": ["T-1"],
        }
    )
    return context


def _visible_content(path: Path) -> tuple[list[str], list[list[list[str]]]]:
    doc = Document(path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    tables = []
    for table in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])
    return paragraphs, tables


def _remove_bookmarks(source: Path, target: Path) -> None:
    with zipfile.ZipFile(source) as source_zip, zipfile.ZipFile(target, "w") as target_zip:
        for info in source_zip.infolist():
            data = source_zip.read(info.filename)
            if info.filename == "word/document.xml":
                root = ET.fromstring(data)
                namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                for element in list(root.iter()):
                    if element.tag in {namespace + "bookmarkStart", namespace + "bookmarkEnd"}:
                        parent = next((candidate for candidate in root.iter() if element in list(candidate)), None)
                        if parent is not None:
                            parent.remove(element)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target_zip.writestr(info, data)


def test_ids_round_trip_and_visible_content_stays_equal() -> None:
    with tempfile.TemporaryDirectory() as temp:
        folder = Path(temp)
        context = _context_with_ids()
        context_path = folder / "context.json"
        context_path.write_text(json.dumps(context, ensure_ascii=False), encoding="utf-8")
        candidate = folder / "candidate.docx"
        control_context = dict(context)
        control_context["blocos"] = [dict(block) for block in context["blocos"]]
        control_context["blocos"][2] = {key: value for key, value in control_context["blocos"][2].items() if key not in {"id", "thesis_ids", "semantic_ids"}}
        control_path = folder / "control.json"
        control_path.write_text(json.dumps(control_context, ensure_ascii=False), encoding="utf-8")
        control = folder / "control.docx"

        for context_file, output in ((context_path, candidate), (control_path, control)):
            result = subprocess.run(
                [sys.executable, str(BUILD), "--context", str(context_file), "--output", str(output)],
                text=True,
                capture_output=True,
            )
            assert result.returncode == 0, result.stdout + result.stderr

        report = verify_docx_semantics(candidate, context)
        assert report["status"] == "PASS"
        assert {"B-1", "T-1"}.issubset(set(report["present"]))
        assert actual_ids_from_docx(candidate)
        assert _visible_content(candidate) == _visible_content(control)


def test_missing_id_blocks_protected_publication_and_preserves_previous_output() -> None:
    with tempfile.TemporaryDirectory() as temp:
        folder = Path(temp)
        context = _context_with_ids()
        context_path = folder / "context.json"
        context_path.write_text(json.dumps(context, ensure_ascii=False), encoding="utf-8")
        generated = folder / "generated.docx"
        stripped = folder / "stripped.docx"
        output = folder / "final.docx"
        result = subprocess.run(
            [sys.executable, str(BUILD), "--context", str(context_path), "--output", str(generated)],
            text=True,
            capture_output=True,
        )
        assert result.returncode == 0, result.stdout + result.stderr
        _remove_bookmarks(generated, stripped)
        report = verify_docx_semantics(stripped, context)
        assert report["status"] == "BLOCK"
        assert "B-1" in report["missing"]
        assert "T-1" in report["missing"]

        output.write_bytes(b"arquivo-anterior")
        publish = subprocess.run(
            [
                sys.executable,
                str(PUBLISH),
                "--input",
                str(stripped),
                "--output",
                str(output),
                "--context",
                str(context_path),
            ],
            text=True,
            capture_output=True,
        )
        assert publish.returncode == 1, publish.stdout + publish.stderr
        assert output.read_bytes() == b"arquivo-anterior"


if __name__ == "__main__":
    test_ids_round_trip_and_visible_content_stays_equal()
    test_missing_id_blocks_protected_publication_and_preserves_previous_output()
    print("[OK] IDs semânticos no DOCX, conteúdo visível e bloqueio protegido passaram")
