#!/usr/bin/env python3
"""Regressões P0 do gerador e do QA estrutural RDAA.

Os testes usam somente arquivos locais e fixtures pequenas. Não chamam modelo,
rede ou serviço externo.
"""

from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

from lxml import etree


PLUGIN_ROOT = Path(__file__).resolve().parents[1]
FIXTURE = PLUGIN_ROOT / "tests" / "fixtures" / "context_happy.json"
GENERATOR = PLUGIN_ROOT / "skills" / "formatar-peca" / "scripts" / "construir_peca.py"
GATE = PLUGIN_ROOT / "skills" / "revisor-rdaa" / "scripts" / "qa_gate.py"
STYLE = PLUGIN_ROOT / "skills" / "revisor-rdaa" / "scripts" / "seguro.py"
STYLE_CHECKER = PLUGIN_ROOT / "skills" / "revisor-rdaa" / "scripts" / "verificar_estilo.py"
NS_URI = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": NS_URI}


def run(command: list[str]) -> subprocess.CompletedProcess[str]:
    # ponytail: sem encoding explícito, o Windows decodifica stdout/stderr com o
    # codepage ANSI do console (cp1252), mesmo quando o processo filho escreve
    # UTF-8 — mensagem com acento vira mojibake ("campo obrigatório" ->
    # "campo obrigatÃ³rio"). errors="replace" evita crash se algo ainda escapar.
    return subprocess.run(
        command, cwd=PLUGIN_ROOT, text=True, capture_output=True,
        encoding="utf-8", errors="replace",
    )


def build_docx(folder: Path) -> Path:
    output = folder / "happy.docx"
    result = run([sys.executable, str(GENERATOR), "--context", str(FIXTURE), "--output", str(output)])
    assert result.returncode == 0, result.stdout + result.stderr
    assert output.is_file()
    return output


def gate(docx: Path, folder: Path) -> subprocess.CompletedProcess[str]:
    result = run([sys.executable, str(GATE), str(docx), "--json", str(folder / f"{docx.stem}.qa.json")])
    return result


def rewrite_zip(source: Path, destination: Path, entry: str, transform) -> None:
    with zipfile.ZipFile(source) as zin, zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            data = zin.read(info.filename)
            if info.filename == entry:
                data = transform(data)
            zout.writestr(info, data)


def parse_xml(data: bytes) -> etree._Element:
    return etree.fromstring(data)


def xml_bytes(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def element_text(element: etree._Element) -> str:
    return "".join(element.itertext())


def mutate_remove_signature_margins(data: bytes) -> bytes:
    root = parse_xml(data)
    for table in root.findall(".//w:tbl", NS):
        if "Wanderley" in element_text(table):
            tbl_pr = table.find("w:tblPr", NS)
            margin = tbl_pr.find("w:tblCellMar", NS) if tbl_pr is not None else None
            if margin is not None:
                tbl_pr.remove(margin)
    return xml_bytes(root)


def mutate_signature_before_close(data: bytes) -> bytes:
    root = parse_xml(data)
    body = root.find("w:body", NS)
    signature = next(child for child in body if child.tag.endswith("}tbl") and "Wanderley" in element_text(child))
    close = next(child for child in body if child.tag.endswith("}p") and "Nestes termos" in element_text(child))
    body.remove(signature)
    body.insert(list(body).index(close), signature)
    return xml_bytes(root)


def mutate_remove_header_respiro(data: bytes) -> bytes:
    root = parse_xml(data)
    paragraphs = root.findall(".//w:p", NS)
    if len(paragraphs) >= 2:
        paragraphs[-1].getparent().remove(paragraphs[-1])
    return xml_bytes(root)


def mutate_insert_body_blank(data: bytes) -> bytes:
    root = parse_xml(data)
    body = root.find("w:body", NS)
    body.insert(0, etree.Element(f"{{{NS_URI}}}p"))
    return xml_bytes(root)


def test_happy_path_and_text_preservation(folder: Path) -> Path:
    docx = build_docx(folder)
    result = gate(docx, folder)
    assert result.returncode == 0, result.stdout + result.stderr
    qa = json.loads((folder / "happy.qa.json").read_text(encoding="utf-8"))
    assert qa["status"] == "PASS"

    from docx import Document

    document = Document(docx)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "A parte autora apresenta esta manifestação" in text
    assert "Nestes termos, aguarda deferimento." in text
    return docx


def assert_mutation(folder: Path, source: Path, name: str, entry: str, transform, expected: str) -> None:
    mutated = folder / f"{name}.docx"
    rewrite_zip(source, mutated, entry, transform)
    result = gate(mutated, folder)
    assert result.returncode != 0, f"mutação {name} passou indevidamente: {result.stdout}"
    assert expected in result.stdout, result.stdout + result.stderr


def test_mutations_suite(folder: Path, source: Path) -> None:
    assert_mutation(folder, source, "no_signature_margin", "word/document.xml", mutate_remove_signature_margins, "margem interna")
    assert_mutation(folder, source, "signature_before_close", "word/document.xml", mutate_signature_before_close, "não está precedida")
    assert_mutation(folder, source, "no_header_respiro", "word/header1.xml", mutate_remove_header_respiro, "parágrafo de respiro")
    assert_mutation(folder, source, "body_blank", "word/document.xml", mutate_insert_body_blank, "corpo iniciado com parágrafo vazio")


def test_invalid_context(folder: Path) -> None:
    invalid = folder / "invalid.json"
    data = json.loads(FIXTURE.read_text(encoding="utf-8"))
    data["blocos"][4].pop("texto")
    invalid.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    output = folder / "invalid.docx"
    result = run([sys.executable, str(GENERATOR), "--context", str(invalid), "--output", str(output)])
    assert result.returncode != 0
    assert "campo obrigatório" in result.stderr
    assert not output.exists()


def test_style_enforcement(folder: Path) -> None:
    from docx import Document

    cases = [
        ("colon.docx", "A tese está demonstrada: a prova documental confirma o fato.", "dois-pontos"),
        ("parenthetical.docx", "A parte autora apresentou o documento (que havia sido solicitado).", "aposto explicativo"),
        ("dash-aside.docx", "A decisão — embora reconheça a pendência — autorizou a penhora.", "travessao proibido"),
        ("single-dash.docx", "A conclusão decorre — o vício está demonstrado.", "travessao proibido"),
        ("semicolon-prose.docx", "O réu não pagou; a autora sofreu prejuízo.", "ponto-e-virgula"),
    ]
    for name, text, expected in cases:
        path = folder / name
        document = Document()
        document.add_paragraph(text)
        document.save(path)
        result = run([sys.executable, str(STYLE_CHECKER), str(path)])
        assert result.returncode == 1, result.stdout + result.stderr
        assert expected in result.stdout, result.stdout + result.stderr

    allowed = folder / "list-marker.docx"
    document = Document()
    document.add_paragraph("A alínea (a) deve ser observada no pedido.")
    document.save(allowed)
    result = run([sys.executable, str(STYLE_CHECKER), str(allowed)])
    assert result.returncode == 0, result.stdout + result.stderr

    # Ementa real de tribunal frequentemente traz dois-pontos ("EMENTA:",
    # "Tema 858:") — citação literal (estilo RDAA Citação) é isenta, a
    # peça não pode reescrever o texto do tribunal pra tirar o dois-pontos.
    citacao_colon = folder / "citacao-colon.docx"
    document = Document()
    from docx.enum.style import WD_STYLE_TYPE
    document.styles.add_style('RDAA Citação', WD_STYLE_TYPE.PARAGRAPH)
    paragraph = document.add_paragraph('EMENTA: RECURSO ESPECIAL. TEMA 858: fixação de tese.')
    paragraph.style = document.styles['RDAA Citação']
    document.save(citacao_colon)
    result = run([sys.executable, str(STYLE_CHECKER), str(citacao_colon)])
    assert result.returncode == 0, result.stdout + result.stderr

    table_colon = folder / "table-colon.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "A tese está demonstrada: a prova documental confirma o fato."
    document.save(table_colon)
    result = run([sys.executable, str(STYLE_CHECKER), str(table_colon)])
    assert result.returncode == 1, result.stdout + result.stderr
    assert "dois-pontos" in result.stdout, result.stdout + result.stderr

    table_allowed = folder / "table-allowed.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "A prova documental confirma o fato."
    document.save(table_allowed)
    result = run([sys.executable, str(STYLE_CHECKER), str(table_allowed)])
    assert result.returncode == 0, result.stdout + result.stderr

    # Ponto-e-vírgula é permitido apenas em parágrafo de lista/alínea.
    semicolon_list = folder / "semicolon-list.docx"
    document = Document()
    from docx.enum.style import WD_STYLE_TYPE
    document.styles.add_style('RDAA Numerado', WD_STYLE_TYPE.PARAGRAPH)
    paragraph = document.add_paragraph("primeiro pedido; segundo pedido; terceiro pedido.")
    paragraph.style = document.styles['RDAA Numerado']
    document.save(semicolon_list)
    result = run([sys.executable, str(STYLE_CHECKER), str(semicolon_list)])
    assert result.returncode == 0, result.stdout + result.stderr

    # Parágrafos consecutivos com a mesma função e estrutura de abertura
    # (mesmo sujeito/verbo) leem mal mesmo com palavras diferentes — ex. real
    # do Ricardo: "esse contexto"/"essa circunstância" são a mesma coisa.
    repeated_opening = folder / "repeated-opening.docx"
    document = Document()
    document.add_paragraph("Esse contexto demonstra que o réu agiu de má-fé durante toda a negociação.")
    document.add_paragraph("Essa circunstância demonstra que o pagamento nunca foi considerado pela perícia.")
    document.save(repeated_opening)
    result = run([sys.executable, str(STYLE_CHECKER), str(repeated_opening)])
    assert result.returncode == 1, result.stdout + result.stderr
    assert "aberturas estruturalmente equivalentes" in result.stdout, result.stdout + result.stderr

    varied_opening = folder / "varied-opening.docx"
    document = Document()
    document.add_paragraph("O comprovante juntado aos autos demonstra que o título já estava quitado.")
    document.add_paragraph("Por sua vez, a perícia reconheceu que houve dupla contagem no cálculo.")
    document.save(varied_opening)
    result = run([sys.executable, str(STYLE_CHECKER), str(varied_opening)])
    assert result.returncode == 0, result.stdout + result.stderr


def test_defensive_openings(folder: Path) -> None:
    from docx import Document

    single = folder / "defensive-opening-single.docx"
    document = Document()
    document.add_paragraph("Não se trata de mera formalidade. O recurso enfrenta o vício identificado na decisão.")
    document.save(single)
    result = run([sys.executable, str(STYLE_CHECKER), str(single)])
    assert result.returncode == 0, result.stdout + result.stderr
    assert "abertura(s) defensiva(s)" in result.stdout

    repeated = folder / "defensive-opening-repeated.docx"
    document = Document()
    document.add_paragraph("Não se pretende rediscutir o mérito. O recurso enfrenta a omissão identificada.")
    document.add_paragraph("Não se pretende rediscutir o mérito. O pedido decorre do vício registrado.")
    document.save(repeated)
    result = run([sys.executable, str(STYLE_CHECKER), str(repeated)])
    assert result.returncode == 1, result.stdout + result.stderr
    assert "Abertura defensiva recorrente" in result.stdout


def test_rollback(folder: Path) -> None:
    stable = folder / "stable.docx"
    new = folder / "new.docx"
    destination = folder / "published.docx"
    backup_dir = folder / "backups"
    stable.write_bytes(b"versao-estavel")
    new.write_bytes(b"versao-nova")
    destination.write_bytes(stable.read_bytes())
    result = run([sys.executable, str(STYLE), "replace", str(new), str(destination), str(backup_dir)])
    assert result.returncode == 0, result.stdout + result.stderr
    backup = next(backup_dir.glob("published.docx.*.bak"))
    assert destination.read_bytes() == b"versao-nova"
    result = run([sys.executable, str(STYLE), "restore", str(backup), str(destination)])
    assert result.returncode == 0, result.stdout + result.stderr
    assert destination.read_bytes() == b"versao-estavel"


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="rdaa-qa-") as tmp:
        folder = Path(tmp)
        source = test_happy_path_and_text_preservation(folder)
        assert_mutation(folder, source, "no_signature_margin", "word/document.xml", mutate_remove_signature_margins, "margem interna")
        assert_mutation(folder, source, "signature_before_close", "word/document.xml", mutate_signature_before_close, "não está precedida")
        assert_mutation(folder, source, "no_header_respiro", "word/header1.xml", mutate_remove_header_respiro, "parágrafo de respiro")
        assert_mutation(folder, source, "body_blank", "word/document.xml", mutate_insert_body_blank, "corpo iniciado com parágrafo vazio")
        test_invalid_context(folder)
        test_style_enforcement(folder)
        test_defensive_openings(folder)
        test_rollback(folder)
    print("[OK] testes P0 de engenharia passaram")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
