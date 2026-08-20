"""Regressões do Visual Law rastreável."""

from __future__ import annotations

import copy
import json
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
BUILD_DIR = ROOT / "skills" / "formatar-peca" / "scripts"
REVIEW_DIR = ROOT / "skills" / "revisor-rdaa" / "scripts"
sys.path.insert(0, str(BUILD_DIR))
sys.path.insert(0, str(REVIEW_DIR))

from construir_peca import construir_peca, validar_contexto  # noqa: E402
from verificar_visual_law import verify_visual_law  # noqa: E402
from estado_rdaa import persist_context  # noqa: E402


def _context() -> dict:
    context = json.loads((ROOT / "tests" / "fixtures" / "context_happy.json").read_text(encoding="utf-8"))
    context["blocos"].append(
        {
            "tipo": "visual",
            "visual_tipo": "timeline",
            "funcao_visual": "Ordenar atos explicitamente fornecidos",
            "texto_pesquisavel": "Evento anonimizado 1 — data informada — ato informado",
            "cabecalho": ["Data", "Evento", "Fonte"],
            "linhas": [["Data informada", "Evento anonimizado 1", "Fonte informada"]],
            "semantic_ids": ["VISUAL-1"],
        }
    )
    return context


def test_visual_law_is_searchable_and_invisible() -> None:
    with tempfile.TemporaryDirectory() as temp:
        context = _context()
        output = Path(temp) / "visual.docx"
        construir_peca(context, output)
        report = verify_visual_law(output, context)
        assert report["status"] == "PASS", report
        from docx import Document
        document = Document(output)
        all_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [paragraph.text for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs]
        )
        assert context["blocos"][-1]["texto_pesquisavel"] in all_text
        visual_table = document.tables[-2]
        hidden_text = "".join(run.text for run in visual_table.rows[0].cells[0].paragraphs[0].runs if run._r.xpath("./w:rPr/w:vanish"))
        assert context["blocos"][-1]["texto_pesquisavel"] in hidden_text


def test_visual_metadata_is_persisted() -> None:
    with tempfile.TemporaryDirectory() as temp:
        context = _context()
        state_dir = Path(temp) / "state"
        persist_context(state_dir, context)
        state = json.loads((state_dir / "matter_state.json").read_text(encoding="utf-8"))
        visual = [item for item in state["semantic_blocks"] if item.get("visual_tipo") == "timeline"]
        assert visual
        assert visual[0]["funcao_visual"] == "Ordenar atos explicitamente fornecidos"


def test_visual_law_requires_function_type_and_search_text() -> None:
    context = _context()
    invalid = copy.deepcopy(context)
    invalid["blocos"][-1].pop("funcao_visual")
    try:
        validar_contexto(invalid)
    except ValueError as exc:
        assert "funcao_visual" in str(exc)
    else:
        raise AssertionError("visual sem função deveria falhar")

    invalid_type = copy.deepcopy(context)
    invalid_type["blocos"][-1]["visual_tipo"] = "decorativo"
    try:
        validar_contexto(invalid_type)
    except ValueError as exc:
        assert "visual_tipo" in str(exc)
    else:
        raise AssertionError("visual com tipo inválido deveria falhar")

    report = verify_visual_law(Path("/tmp/arquivo-inexistente.docx"), invalid_type)
    assert report["status"] == "BLOCK"


def main() -> int:
    test_visual_law_is_searchable_and_invisible()
    test_visual_metadata_is_persisted()
    test_visual_law_requires_function_type_and_search_text()
    print("[OK] Visual Law tipado, pesquisável e invisível passou")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
